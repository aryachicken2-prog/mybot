try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    DOCX_AVAILABLE = False
import tempfile
from telethon import events, Button
import importlib
import platform
from utils import (
    CANCEL_BUTTON, BACK_BUTTON, paginate_buttons,
    set_user_state, get_user_state, get_user_data, clear_user_state,
    is_admin as utils_is_admin, CHANNEL_USERNAME, get_setting, set_setting, is_safe_upload_path
)
from database import DB_NAME, OWNER_ID
import sqlite3
try:
    import pandas as pd
    PD_AVAILABLE = True
except Exception:
    pd = None
    PD_AVAILABLE = False
try:
    import xlsxwriter  # type: ignore[reportMissingImports]
except Exception:
    xlsxwriter = None
from io import BytesIO
import json
import os
import time
import random
import asyncio

from io import BytesIO as _BytesIO


def apply_persian_xlsx_style(buf: _BytesIO) -> _BytesIO:
    """If openpyxl is available, load the workbook from BytesIO, set right-to-left,
    set font to 'B Nazanin' and right alignment for cells, adjust column widths,
    and return a new BytesIO. If anything fails, return original buffer.
    """
    try:
        buf.seek(0)
        from openpyxl import load_workbook
        from openpyxl.styles import Font, Alignment
        from openpyxl.utils import get_column_letter, column_index_from_string
    except Exception:
        try:
            buf.seek(0)
        except Exception:
            pass
        return buf

    try:
        # load workbook from BytesIO and apply RTL/font/alignment/widths
        wb = load_workbook(filename=buf)
        nazanin = Font(name='B Nazanin')
        align_r = Alignment(horizontal='right', vertical='top', wrap_text=True)
        for ws in wb.worksheets:
            try:
                ws.sheet_view.rightToLeft = True
            except Exception:
                pass
            # apply font & alignment to all cells and compute column widths
            max_widths = {}
            for row in ws.iter_rows():
                for cell in row:
                    try:
                        if cell.value is not None:
                            s = str(cell.value)
                        else:
                            s = ''
                        cell.font = nazanin
                        cell.alignment = align_r
                        col = cell.column
                        # openpyxl column might be int or string; normalize
                        try:
                            idx = int(col)
                        except Exception:
                            try:
                                idx = column_index_from_string(col)
                            except Exception:
                                continue
                        prev = max_widths.get(idx, 0)
                        if len(s) > prev:
                            max_widths[idx] = len(s)
                    except Exception:
                        continue
            # set widths
            for idx, width in max_widths.items():
                try:
                    letter = get_column_letter(idx)
                    ws.column_dimensions[letter].width = min(50, max(10, int(width * 1.2)))
                except Exception:
                    pass

        out = _BytesIO()
        wb.save(out)
        out.seek(0)
        return out
    except Exception:
        try:
            buf.seek(0)
        except Exception:
            pass
        return buf


def apply_persian_docx_style(doc):
    """Apply right-to-left alignment and B Nazanin font to a python-docx Document.
    This attempts to set paragraph alignment to RIGHT and apply font name to runs.
    It is best-effort and will silently continue on failure.
    """
    try:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.oxml.ns import qn
    except Exception:
        return doc
    try:
        # Paragraph-level
        for paragraph in list(doc.paragraphs):
            try:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            except Exception:
                pass
            for run in paragraph.runs:
                try:
                    run.font.name = 'B Nazanin'
                    # also set eastAsia font for some Word renderers
                    try:
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                    except Exception:
                        pass
                except Exception:
                    pass

        # Table cells
        for table in list(doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        try:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                        except Exception:
                            pass
                        for run in paragraph.runs:
                            try:
                                run.font.name = 'B Nazanin'
                                try:
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'B Nazanin')
                                except Exception:
                                    pass
                            except Exception:
                                pass
    except Exception:
        return doc
    return doc


def safe_df_from_rows(rows, headers):
    """Create a pandas DataFrame from rows and headers in a tolerant way.
    If rows have different lengths, pad them with None/empty strings.
    """
    try:
        import pandas as _pd
    except Exception:
        return None
    norm = []
    for r in rows:
        if not isinstance(r, (list, tuple)):
            r = [r]
        # pad or truncate to headers length
        row = list(r)[:len(headers)] + [None] * max(0, len(headers) - len(r))
        norm.append(row)
    try:
        return _pd.DataFrame(norm, columns=headers)
    except Exception:
        try:
            return _pd.DataFrame(norm)
        except Exception:
            return None

def setup_admin_handlers(client, user_states):
    @client.on(events.CallbackQuery)
    async def admin_callback_handler(event):
        data = event.data.decode('utf-8')
        user_id = event.sender_id

        if data.startswith("admin_message_members_"):
            # انتخاب گروه هدف (approved/rejected/all)
            target = data.split("_")[-1]
            if target == "approved":
                status = "approved"
            elif target == "rejected":
                status = "rejected"
            else:
                status = None
            set_user_state(user_states, user_id, "admin_waiting_members_message", {"target_status": status})
            await event.edit("✏️ لطفاً پیام مورد نظر را بنویسید و ارسال کنید:", buttons=CANCEL_BUTTON)
            return

    @client.on(events.NewMessage)
    async def admin_message_handler(event):
        user_id = event.sender_id
        state = get_user_state(user_states, user_id)
        if state == "admin_waiting_user_help_text":
            text = event.message.text or ""
            if not text.strip():
                await event.reply("❌ متن راهنما نمی‌تواند خالی باشد. لطفاً دوباره ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            set_setting('user_help_text', text.strip())
            clear_user_state(user_states, user_id)
            await event.reply("✅ متن راهنمای ربات با موفقیت ذخیره شد.", buttons=get_admin_main_menu())
            return
        if state == "admin_waiting_membership_desc":
            text = event.message.text or ""
            if not text.strip():
                await event.reply("❌ متن توضیحات عضویت نمی‌تواند خالی باشد. لطفاً دوباره ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            set_setting('membership_description', text.strip())
            clear_user_state(user_states, user_id)
            await event.reply("✅ متن توضیحات عضویت با موفقیت ذخیره شد.", buttons=get_admin_main_menu())
            return
        if state == "admin_waiting_members_message":
            text = event.message.text or ""
            if not text.strip():
                await event.reply("❌ پیام نمی‌تواند خالی باشد. لطفاً دوباره ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            target_status = get_user_data(user_states, user_id).get("target_status")
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            if target_status:
                c.execute("SELECT user_id FROM memberships WHERE status=?", (target_status,))
            else:
                c.execute("SELECT user_id FROM memberships")
            user_ids = [row[0] for row in c.fetchall()]
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply(f"⏳ ارسال پیام به {len(user_ids)} کاربر در حال انجام است...", buttons=get_admin_main_menu())
            sent, failed = 0, 0
            for uid in user_ids:
                try:
                    await client.send_message(uid, text)
                    sent += 1
                except Exception:
                    failed += 1
            await event.reply(f"✅ پیام به {sent} کاربر ارسال شد. {failed} مورد ناموفق.", buttons=get_admin_main_menu())
            return
        if state == "admin_waiting_target_id":
            # admin should have sent a numeric id
            text = getattr(event.message, 'text', '') or ''
            txt = text.strip()
            if not txt.isdigit():
                await event.reply("❌ آیدی نامعتبر است. لطفا فقط عدد را ارسال کنید (برای خروج /cancel).", buttons=CANCEL_BUTTON)
                return
            target_id = int(txt)
            # store target and prompt for message content
            data = get_user_data(user_states, user_id) or {}
            data['target_id'] = target_id
            set_user_state(user_states, user_id, "admin_waiting_send_to_id_content", data)
            await event.reply(f"✏️ حالا پیام یا فایل را برای ارسال به {target_id} ارسال کنید:\n(متن، عکس یا ویدیو پذیرفته می‌شود)", buttons=CANCEL_BUTTON)
            return
        if state == "admin_waiting_send_to_id_content":
            # send text or media to stored target id
            data = get_user_data(user_states, user_id) or {}
            target_id = data.get('target_id')
            if not target_id:
                clear_user_state(user_states, user_id)
                await event.reply("❌ آیدی مخاطب یافت نشد. عملیات لغو شد.", buttons=get_admin_main_menu())
                return
            sent, failed = 0, 0
            # support text messages
            try:
                # if message contains media, forward the media or send the file
                if getattr(event.message, 'media', None) is not None:
                    # try to forward the exact message (keeps caption)
                    try:
                        await client.send_file(target_id, event.message.media, caption=(event.message.message or ''))
                        sent = 1
                    except Exception:
                        # fallback: try to forward message (keeps sender), may fail if can't forward
                        try:
                            await event.message.forward_to(target_id)
                            sent = 1
                        except Exception:
                            failed = 1
                else:
                    text = event.message.text or ''
                    if not (text and text.strip()):
                        await event.reply("❌ پیام خالی است. لطفاً متن یا مدیا ارسال کنید.", buttons=CANCEL_BUTTON)
                        return
                    try:
                        await client.send_message(int(target_id), text)
                        sent = 1
                    except Exception:
                        failed = 1
            except Exception:
                failed = 1

            clear_user_state(user_states, user_id)
            if sent:
                await event.reply(f"✅ پیام به {target_id} ارسال شد.", buttons=get_admin_main_menu())
            else:
                await event.reply(f"❌ ارسال به {target_id} ناموفق بود.", buttons=get_admin_main_menu())
            return

    @client.on(events.CallbackQuery)
    async def admin_callback_handler(event):
        data = event.data.decode('utf-8')
        user_id = event.sender_id

        # DEBUG: log incoming callback and current state for troubleshooting
        try:
            from log_helper import console_log
            try:
                st = get_user_state(user_states, user_id)
            except Exception:
                st = None
            console_log(f"admin_callback: user={user_id} data={data} state={st}")
        except Exception:
            try:
                print(f"[DEBUG] admin_callback: user={user_id} data={data} state={get_user_state(user_states, user_id)}")
            except Exception:
                print(f"[DEBUG] admin_callback: user={user_id} data={data}")

        if not utils_is_admin(DB_NAME, user_id):
            return

        if data == "cancel":
            clear_user_state(user_states, user_id)
            try:
                await event.edit("✅ عملیات لغو شد.", buttons=get_admin_main_menu())
            except:
                await event.respond("✅ عملیات لغو شد.", buttons=get_admin_main_menu())
            return

        if data == "main_menu":
            clear_user_state(user_states, user_id)
            await event.edit("منوی اصلی ادمین:", buttons=get_admin_main_menu())
            return

        if data == "admin_events":
            buttons = [
                [Button.inline("➕ ثبت رویداد جدید", b"admin_new_event_step1"), Button.inline("⚙️ مدیریت رویدادها", b"admin_manage_events_0")],
                [Button.inline("🔙 بازگشت", b"main_menu")]
            ]
            await event.edit("📅 بخش مدیریت رویدادها:", buttons=buttons)

        # Handle cost type selection during new event creation
        elif data == "cost_free":
            # Ask admin whether certificate of attendance will be issued for free events
            buttons = [
                [Button.inline("✅ بله — گواهی صادر می‌شود", b"cert_yes")],
                [Button.inline("❌ خیر — گواهی صادر نمی‌شود", b"cert_no")],
                [Button.inline("🔙 بازگشت", b"main_menu")]
            ]
            await event.edit("این رویداد رایگان است. آیا گواهی حضور برای شرکت‌کنندگان صادر خواهد شد؟", buttons=buttons)
            return
            # compute labels for other notification toggles
            nmemb = "✅ روشن" if get_setting("notify_new_membership","1") == "1" else "❌ خاموش"
            nidea = "✅ روشن" if get_setting("notify_new_idea","1") == "1" else "❌ خاموش"
            ncollab = "✅ روشن" if get_setting("notify_new_collab","1") == "1" else "❌ خاموش"
            ndonate = "✅ روشن" if get_setting("notify_new_donation","1") == "1" else "❌ خاموش"

        elif data == "cert_no":
            # Proceed as free event without certificate fields
            # mark cost_type free and continue to poster upload step
            try:
                # set user state data
                data = get_user_data(user_states, user_id) or {}
                data["cost_type"] = "free"
                # DEBUG: log who sets the poster state
                try:
                    from log_helper import console_log
                    console_log(f"SET_STATE: user={user_id} path=cert_no -> admin_new_event_poster data_keys={list(data.keys())}")
                except Exception:
                    try:
                        print(f"[DEBUG] SET_STATE: user={user_id} path=cert_no -> admin_new_event_poster data_keys={list(data.keys())}")
                    except Exception:
                        pass
                set_user_state(user_states, user_id, "admin_new_event_poster", data)
                await event.edit("🏷️ رویداد به عنوان رایگان ثبت خواهد شد. اکنون لطفا پوستر رویداد را ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            except Exception:
                await event.answer("خطا در پردازش، لطفاً دوباره تلاش کنید.", alert=True)
                return

        elif data == "cert_yes":
            # Ask whether cert fee differs for student / non-student
            try:
                data = get_user_data(user_states, user_id) or {}
                data["cost_type"] = "free"
                set_user_state(user_states, user_id, "admin_new_event_cert_diff", data)
                buttons = [
                    [Button.inline("✅ بله — متفاوت است", b"cert_diff_yes")],
                    [Button.inline("❌ خیر — یکسان است", b"cert_diff_no")],
                    [Button.inline("🔙 بازگشت", b"main_menu")]
                ]
                await event.edit("آیا هزینه صدور گواهی برای دانشجو و غیر دانشجو متفاوت است؟", buttons=buttons)
                return
            except Exception:
                await event.answer("خطا در پردازش، لطفاً دوباره تلاش کنید.", alert=True)
                return

        elif data == "cert_diff_no":
            # collect a single cert fee (same for all)
            try:
                data = get_user_data(user_states, user_id) or {}
                data["cost_type"] = "free"
                set_user_state(user_states, user_id, "admin_new_event_cert_fee", data)
                await event.edit("📥 لطفا میزان هزینه صدور گواهی برای هر شرکت‌کننده را به تومان وارد کنید (عدد):", buttons=CANCEL_BUTTON)
                return
            except Exception:
                await event.answer("خطا در پردازش، لطفاً دوباره تلاش کنید.", alert=True)
                return

        elif data == "cert_diff_yes":
            # collect student cert fee first, then non-student
            try:
                data = get_user_data(user_states, user_id) or {}
                data["cost_type"] = "free"
                set_user_state(user_states, user_id, "admin_new_event_cert_fee_student", data)
                await event.edit("📥 لطفا میزان هزینه صدور گواهی برای دانشجو را به تومان وارد کنید (عدد):", buttons=CANCEL_BUTTON)
                return
            except Exception:
                await event.answer("خطا در پردازش، لطفاً دوباره تلاش کنید.", alert=True)
                return

        elif data == "cost_fixed":
            data = get_user_data(user_states, user_id) or {}
            data["cost_type"] = "fixed"
            set_user_state(user_states, user_id, "admin_new_event_cost_amount", data)
            await event.edit("💰 لطفا مبلغ ثابت رویداد را به تومان وارد کنید (فقط عدد):", buttons=CANCEL_BUTTON)

        elif data == "cost_variable":
            data = get_user_data(user_states, user_id) or {}
            data["cost_type"] = "variable"
            set_user_state(user_states, user_id, "admin_new_event_cost_amount", data)
            await event.edit("🎓 لطفا هزینه برای دانشجو را به تومان وارد کنید (عدد):", buttons=CANCEL_BUTTON)

        elif data == "admin_maintenance":
            buttons = [
                [Button.inline("🧽 پاکسازی ثبت‌نام‌های در انتظار", b"maint_clear_pending_regs"), Button.inline("🗑️ حذف تیکت‌های بسته", b"maint_clear_closed_tickets")],
                [Button.inline("🧾 حذف رسیدهای بدون ثبت‌نام", b"maint_clear_orphan_receipts"), Button.inline("🧹 پاکسازی فایل‌های اضافی سرور", b"maint_clear_stray_files")],
                [Button.inline("🗂️ حذف داده‌های یک رویداد", b"maint_purge_event_select")],
                [Button.inline("🧺 حذف فقط فایل‌های یک رویداد", b"maint_purge_event_files_select")],
                [Button.inline("🏠 بازگشت", b"main_menu")]
            ]
            await event.edit("🧹 بخش نگهداری و پاکسازی:", buttons=buttons)
        elif data == "maint_purge_event_select":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events ORDER BY id DESC")
            evs = c.fetchall()
            conn.close()
            if not evs:
                await event.answer("رویدادی وجود ندارد.", alert=True)
                return
            buttons = []
            for eid, title in evs[:50]:
                buttons.append([Button.inline(f"🗑️ {title}", f"maint_purge_event_{eid}")])
            buttons.append([Button.inline("🔙 بازگشت", b"admin_maintenance")])
            await event.edit("لطفاً رویداد موردنظر برای حذف کامل داده‌ها را انتخاب کنید:", buttons=buttons)

        elif data == "maint_purge_event_files_select":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events ORDER BY id DESC")
            evs = c.fetchall()
            conn.close()
            if not evs:
                await event.answer("رویدادی وجود ندارد.", alert=True)
                return
            buttons = []
            for eid, title in evs[:50]:
                buttons.append([Button.inline(f"🧺 فایل‌های {title}", f"maint_purge_event_files_{eid}")])
            buttons.append([Button.inline("🔙 بازگشت", b"admin_maintenance")])
            await event.edit("لطفاً رویداد موردنظر برای حذف فقط فایل‌ها را انتخاب کنید:", buttons=buttons)

        elif data == "admin_settings":
            # allow any admin to view/change these settings
            if not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            nreg = "✅ روشن" if get_setting("notify_new_registration","1") == "1" else "❌ خاموش"
            ntick = "✅ روشن" if get_setting("notify_new_ticket","1") == "1" else "❌ خاموش"
            single_reg = "✅ روشن" if get_setting("single_registration_per_user","1") == "1" else "❌ خاموش"
            # ensure membership/idea/collab/donation labels are defined to avoid UnboundLocalError
            nmemb = "✅ روشن" if get_setting("notify_new_membership","1") == "1" else "❌ خاموش"
            nidea = "✅ روشن" if get_setting("notify_new_idea","1") == "1" else "❌ خاموش"
            ncollab = "✅ روشن" if get_setting("notify_new_collab","1") == "1" else "❌ خاموش"
            ndonate = "✅ روشن" if get_setting("notify_new_donation","1") == "1" else "❌ خاموش"
            try:
                donate_card = get_setting('donation_card_number', '—')
            except Exception:
                donate_card = '—'
            try:
                donate_desc_preview = get_setting('donation_description', '')
            except Exception:
                donate_desc_preview = ''
            donate_label = donate_card if donate_card else '—'
            buttons = [
                [Button.inline(f"ثبت‌نام جدید: {nreg}", b"toggle_notify_reg")],
                [Button.inline(f"تیکت جدید: {ntick}", b"toggle_notify_ticket")],
                [Button.inline(f"درخواست عضویت: {nmemb}", b"toggle_notify_membership")],
                [Button.inline(f"ایده جدید: {nidea}", b"toggle_notify_ideas")],
                [Button.inline(f"درخواست همکاری: {ncollab}", b"toggle_notify_collabs")],
                [Button.inline(f"حمایت جدید: {ndonate}", b"toggle_notify_donations")],
                [Button.inline(f"🔒 یک ثبت‌نام/کاربر: {single_reg}", b"toggle_single_reg")],
                [Button.inline("✏️ ویرایش", b"admin_edit_menu")],
                [Button.inline("🏠 بازگشت", b"main_menu")]
            ]
            await event.edit("⚙️ تنظیمات اعلان‌ها:", buttons=buttons)
        # admin_diag handler removed
        elif data == "admin_edit_menu":
            # New edit menu: move donation/membership editing here
            if not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            donate_card = get_setting('donation_card_number', '—')
            donate_holder = get_setting('donation_card_holder', '')
            donate_desc_preview = get_setting('donation_description', '')
            membership_desc_preview = get_setting('membership_description', '')
            buttons = [
                [Button.inline(f"💳 شماره کارت حمایت: {donate_card}", b"admin_set_donation_card")],
                [Button.inline(f"🧾 صاحب کارت: {donate_holder or '—'}", b"admin_set_donation_holder")],
                [Button.inline("📝 توضیحات حمایت (ویرایش)", b"admin_set_donation_desc")],
                [Button.inline("📝 توضیحات عضویت (ویرایش)", b"admin_edit_membership_desc")],
                [Button.inline("🔙 بازگشت", b"main_menu")]
            ]
            await event.edit("✏️ منوی ویرایش:", buttons=buttons)
        elif data == "admin_export_members_excel":
            await event.edit("⏳ در حال تهیه فایل اکسل عضویت...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, full_name, major, entry_year, student_number, national_id, phone, telegram_username, student_card_file, status, created_at FROM memberships ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            headers = ["شماره", "آیدی کاربر", "نام و نام خانوادگی", "رشته", "سال ورود", "شماره دانشجویی", "کد ملی", "تلفن", "آیدی تلگرام", "فایل کارت دانشجویی", "وضعیت", "تاریخ"]
            # Prefer pandas -> xlsx. If pandas missing, try openpyxl or xlsxwriter. Otherwise CSV fallback.
            if PD_AVAILABLE:
                try:
                    df = safe_df_from_rows(rows, headers)
                    bio = BytesIO()
                    df.to_excel(bio, index=False, engine='openpyxl')
                    bio.seek(0)
                    bio = apply_persian_xlsx_style(bio)
                    try:
                        bio.name = 'members.xlsx'
                    except Exception:
                        pass
                    await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل عضویت", force_document=True)
                    await event.edit("✅ فایل اکسل عضویت ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception as e:
                    try:
                        from log_helper import console_log
                        import traceback as _tb
                        console_log(f"export_members_excel pandas error: {_tb.format_exc()}")
                    except Exception:
                        import traceback as _tb
                        print("export_members_excel pandas error:")
                        print(_tb.format_exc())

            # Try openpyxl directly
            try:
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.append(headers)
                for r in rows:
                    ws.append(list(r))
                bio = BytesIO()
                wb.save(bio)
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'members.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل عضویت", force_document=True)
                await event.edit("✅ فایل اکسل عضویت ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception as e:
                try:
                    from log_helper import console_log
                    import traceback as _tb
                    console_log(f"export_members_excel openpyxl error: {_tb.format_exc()}")
                except Exception:
                    import traceback as _tb
                    print("export_members_excel openpyxl error:")
                    print(_tb.format_exc())

            # Try xlsxwriter
            try:
                import xlsxwriter as _xlsxwriter
                bio = BytesIO()
                workbook = _xlsxwriter.Workbook(bio)
                ws = workbook.add_worksheet()
                for c_idx, h in enumerate(headers):
                    ws.write(0, c_idx, h)
                for r_idx, r in enumerate(rows, start=1):
                    for c_idx, val in enumerate(r):
                        ws.write(r_idx, c_idx, val)
                workbook.close()
                bio.seek(0)
                # xlsxwriter output is a valid xlsx; attempt to style via openpyxl
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'members.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل عضویت", force_document=True)
                await event.edit("✅ فایل اکسل عضویت ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception as e:
                try:
                    from log_helper import console_log
                    import traceback as _tb
                    console_log(f"export_members_excel xlsxwriter error: {_tb.format_exc()}")
                except Exception:
                    import traceback as _tb
                    print("export_members_excel xlsxwriter error:")
                    print(_tb.format_exc())

            # CSV fallback
            try:
                import csv
                with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, encoding='utf-8', newline='') as tmp:
                    writer = csv.writer(tmp)
                    writer.writerow(headers)
                    for r in rows:
                        writer.writerow(r)
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📥 فایل CSV عضویت (fallback)", force_document=True)
                await event.edit("✅ فایل CSV عضویت ارسال شد (fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return
        elif data == "admin_export_ideas_word":
            await event.edit("⏳ در حال تهیه فایل خروجی ایده‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, title, description, status, created_at FROM ideas ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            # Prefer docx, otherwise fallback to plain text
            if DOCX_AVAILABLE:
                try:
                    doc = Document()
                    doc.add_heading("ایده‌های ثبت شده", 0)
                    table = doc.add_table(rows=1, cols=6)
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'شماره'
                    hdr_cells[1].text = 'کاربر'
                    hdr_cells[2].text = 'عنوان'
                    hdr_cells[3].text = 'توضیحات'
                    hdr_cells[4].text = 'وضعیت'
                    hdr_cells[5].text = 'تاریخ'
                    for r in rows:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(r[0])
                        row_cells[1].text = str(r[1])
                        row_cells[2].text = str(r[2] or '')
                        row_cells[3].text = str(r[3] or '')
                        row_cells[4].text = str(r[4] or '')
                        row_cells[5].text = str(r[5] or '')
                    try:
                        doc = apply_persian_docx_style(doc)
                    except Exception:
                        pass
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        doc.save(tmp.name)
                        tmp.flush()
                        await client.send_file(event.chat_id, tmp.name, caption="📄 فایل ورد ایده‌ها", force_document=True)
                    await event.edit("✅ فایل ورد ایده‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            # TXT fallback
            try:
                with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', delete=False, encoding='utf-8') as tmp:
                    for row in rows:
                        tmp.write(f"شماره: {row[0]}\n")
                        tmp.write(f"کاربر: {row[1]}\n")
                        tmp.write(f"عنوان: {row[2]}\n")
                        tmp.write(f"توضیحات: {row[3]}\n")
                        tmp.write(f"وضعیت: {row[4]}\n")
                        tmp.write(f"تاریخ: {row[5]}\n")
                        tmp.write("-\n")
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📄 خروجی ایده‌ها (TXT fallback)", force_document=True)
                await event.edit("✅ خروجی ایده‌ها ارسال شد (TXT fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return
        elif data == "admin_export_collabs_word":
            await event.edit("⏳ در حال تهیه فایل خروجی همکاری‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, full_name, organization, proposal, status, created_at FROM collaborations ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            if DOCX_AVAILABLE:
                try:
                    doc = Document()
                    doc.add_heading("درخواست‌های همکاری", 0)
                    table = doc.add_table(rows=1, cols=7)
                    hdr = table.rows[0].cells
                    hdr[0].text = 'شماره'
                    hdr[1].text = 'کاربر'
                    hdr[2].text = 'نام'
                    hdr[3].text = 'سازمان'
                    hdr[4].text = 'پیشنهاد'
                    hdr[5].text = 'وضعیت'
                    hdr[6].text = 'تاریخ'
                    for r in rows:
                        cells = table.add_row().cells
                        cells[0].text = str(r[0])
                        cells[1].text = str(r[1])
                        cells[2].text = str(r[2] or '')
                        cells[3].text = str(r[3] or '')
                        cells[4].text = str(r[4] or '')
                        cells[5].text = str(r[5] or '')
                        cells[6].text = str(r[6] or '')
                    try:
                        doc = apply_persian_docx_style(doc)
                    except Exception:
                        pass
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        doc.save(tmp.name)
                        tmp.flush()
                        await client.send_file(event.chat_id, tmp.name, caption="📄 فایل ورد همکاری‌ها", force_document=True)
                    await event.edit("✅ فایل ورد همکاری‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', delete=False, encoding='utf-8') as tmp:
                    for row in rows:
                        tmp.write(f"شماره: {row[0]}\n")
                        tmp.write(f"کاربر: {row[1]}\n")
                        tmp.write(f"نام: {row[2]}\n")
                        tmp.write(f"سازمان: {row[3]}\n")
                        tmp.write(f"پیشنهاد: {row[4]}\n")
                        tmp.write(f"وضعیت: {row[5]}\n")
                        tmp.write(f"تاریخ: {row[6]}\n")
                        tmp.write("-\n")
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📄 خروجی همکاری‌ها (TXT fallback)", force_document=True)
                await event.edit("✅ خروجی همکاری‌ها ارسال شد (TXT fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return
        elif data == "admin_export_donations_word":
            await event.edit("⏳ در حال تهیه فایل خروجی حمایت‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, amount, currency, status, created_at FROM donations ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            if DOCX_AVAILABLE:
                try:
                    doc = Document()
                    doc.add_heading("حمایت‌های ثبت شده", 0)
                    table = doc.add_table(rows=1, cols=6)
                    hdr = table.rows[0].cells
                    hdr[0].text = 'شماره'
                    hdr[1].text = 'کاربر'
                    hdr[2].text = 'مبلغ'
                    hdr[3].text = 'واحد'
                    hdr[4].text = 'وضعیت'
                    hdr[5].text = 'تاریخ'
                    for r in rows:
                        cells = table.add_row().cells
                        cells[0].text = str(r[0])
                        cells[1].text = str(r[1])
                        cells[2].text = str(r[2])
                        cells[3].text = str(r[3])
                        cells[4].text = str(r[4] or '')
                        cells[5].text = str(r[5] or '')
                    try:
                        doc = apply_persian_docx_style(doc)
                    except Exception:
                        pass
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        doc.save(tmp.name)
                        tmp.flush()
                        await client.send_file(event.chat_id, tmp.name, caption="📄 فایل ورد حمایت‌ها", force_document=True)
                    await event.edit("✅ فایل ورد حمایت‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', delete=False, encoding='utf-8') as tmp:
                    for row in rows:
                        tmp.write(f"شماره: {row[0]}\n")
                        tmp.write(f"کاربر: {row[1]}\n")
                        tmp.write(f"مبلغ: {row[2]} {row[3]}\n")
                        tmp.write(f"وضعیت: {row[4]}\n")
                        tmp.write(f"تاریخ: {row[5]}\n")
                        tmp.write("-\n")
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📄 خروجی حمایت‌ها (TXT fallback)", force_document=True)
                await event.edit("✅ خروجی حمایت‌ها ارسال شد (TXT fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return

        elif data == "admin_export_tickets_word":
            await event.edit("⏳ در حال تهیه فایل خروجی تیکت‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, message, status, created_at FROM tickets ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            if DOCX_AVAILABLE:
                try:
                    doc = Document()
                    doc.add_heading("تیکت‌های پشتیبانی", 0)
                    table = doc.add_table(rows=1, cols=5)
                    hdr = table.rows[0].cells
                    hdr[0].text = 'شماره'
                    hdr[1].text = 'کاربر'
                    hdr[2].text = 'پیام'
                    hdr[3].text = 'وضعیت'
                    hdr[4].text = 'تاریخ'
                    for r in rows:
                        cells = table.add_row().cells
                        cells[0].text = str(r[0])
                        cells[1].text = str(r[1])
                        cells[2].text = str(r[2] or '')
                        cells[3].text = str(r[3] or '')
                        cells[4].text = str(r[4] or '')
                    try:
                        doc = apply_persian_docx_style(doc)
                    except Exception:
                        pass
                    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                        doc.save(tmp.name)
                        tmp.flush()
                        await client.send_file(event.chat_id, tmp.name, caption="📄 فایل ورد تیکت‌ها", force_document=True)
                    await event.edit("✅ فایل ورد تیکت‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                with tempfile.NamedTemporaryFile(suffix='.txt', mode='w', delete=False, encoding='utf-8') as tmp:
                    for row in rows:
                        tmp.write(f"شماره: {row[0]}\n")
                        tmp.write(f"کاربر: {row[1]}\n")
                        tmp.write(f"پیام: {row[2]}\n")
                        tmp.write(f"وضعیت: {row[3]}\n")
                        tmp.write(f"تاریخ: {row[4]}\n")
                        tmp.write("-\n")
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📄 خروجی تیکت‌ها (TXT fallback)", force_document=True)
                await event.edit("✅ خروجی تیکت‌ها ارسال شد (TXT fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return

        elif data == "admin_export_ideas_excel":
            await event.edit("⏳ در حال تهیه فایل اکسل ایده‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, title, description, status, created_at FROM ideas ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            headers = ["شماره", "کاربر", "عنوان", "توضیحات", "وضعیت", "تاریخ"]
            # Prefer pandas
            if PD_AVAILABLE:
                try:
                    df = safe_df_from_rows(rows, headers)
                    bio = BytesIO()
                    df.to_excel(bio, index=False, engine='openpyxl')
                    bio.seek(0)
                    bio = apply_persian_xlsx_style(bio)
                    try:
                        bio.name = 'ideas.xlsx'
                    except Exception:
                        pass
                    await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل ایده‌ها", force_document=True)
                    await event.edit("✅ فایل اکسل ایده‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.append(headers)
                for r in rows:
                    ws.append(list(r))
                bio = BytesIO()
                wb.save(bio)
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'ideas.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل ایده‌ها", force_document=True)
                await event.edit("✅ فایل اکسل ایده‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import xlsxwriter as _xlsxwriter
                bio = BytesIO()
                workbook = _xlsxwriter.Workbook(bio)
                ws = workbook.add_worksheet()
                for c_idx, h in enumerate(headers):
                    ws.write(0, c_idx, h)
                for r_idx, r in enumerate(rows, start=1):
                    for c_idx, val in enumerate(r):
                        ws.write(r_idx, c_idx, val)
                workbook.close()
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'ideas.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل ایده‌ها", force_document=True)
                await event.edit("✅ فایل اکسل ایده‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import csv
                with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, encoding='utf-8', newline='') as tmp:
                    writer = csv.writer(tmp)
                    writer.writerow(headers)
                    for r in rows:
                        writer.writerow(r)
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📥 فایل CSV ایده‌ها (fallback)", force_document=True)
                await event.edit("✅ فایل CSV ایده‌ها ارسال شد (fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return

        elif data == "admin_export_collabs_excel":
            await event.edit("⏳ در حال تهیه فایل اکسل همکاری‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, full_name, organization, proposal, status, created_at FROM collaborations ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            headers = ["شماره", "کاربر", "نام", "سازمان", "پیشنهاد", "وضعیت", "تاریخ"]
            if PD_AVAILABLE:
                try:
                    df = safe_df_from_rows(rows, headers)
                    bio = BytesIO()
                    df.to_excel(bio, index=False, engine='openpyxl')
                    bio.seek(0)
                    bio = apply_persian_xlsx_style(bio)
                    try:
                        bio.name = 'collabs.xlsx'
                    except Exception:
                        pass
                    await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل همکاری‌ها", force_document=True)
                    await event.edit("✅ فایل اکسل همکاری‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.append(headers)
                for r in rows:
                    ws.append(list(r))
                bio = BytesIO()
                wb.save(bio)
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'collabs.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل همکاری‌ها", force_document=True)
                await event.edit("✅ فایل اکسل همکاری‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import xlsxwriter as _xlsxwriter
                bio = BytesIO()
                workbook = _xlsxwriter.Workbook(bio)
                ws = workbook.add_worksheet()
                for c_idx, h in enumerate(headers):
                    ws.write(0, c_idx, h)
                for r_idx, r in enumerate(rows, start=1):
                    for c_idx, val in enumerate(r):
                        ws.write(r_idx, c_idx, val)
                workbook.close()
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'collabs.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل همکاری‌ها", force_document=True)
                await event.edit("✅ فایل اکسل همکاری‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import csv
                with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, encoding='utf-8', newline='') as tmp:
                    writer = csv.writer(tmp)
                    writer.writerow(headers)
                    for r in rows:
                        writer.writerow(r)
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📥 فایل CSV همکاری‌ها (fallback)", force_document=True)
                await event.edit("✅ فایل CSV همکاری‌ها ارسال شد (fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return

        elif data == "admin_export_donations_excel":
            await event.edit("⏳ در حال تهیه فایل اکسل حمایت‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, amount, currency, status, created_at FROM donations ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            headers = ["شماره", "کاربر", "مبلغ", "واحد", "وضعیت", "تاریخ"]
            if PD_AVAILABLE:
                try:
                    df = safe_df_from_rows(rows, headers)
                    bio = BytesIO()
                    df.to_excel(bio, index=False, engine='openpyxl')
                    bio.seek(0)
                    bio = apply_persian_xlsx_style(bio)
                    try:
                        bio.name = 'donations.xlsx'
                    except Exception:
                        pass
                    await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل حمایت‌ها", force_document=True)
                    await event.edit("✅ فایل اکسل حمایت‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.append(headers)
                for r in rows:
                    ws.append(list(r))
                bio = BytesIO()
                wb.save(bio)
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'donations.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل حمایت‌ها", force_document=True)
                await event.edit("✅ فایل اکسل حمایت‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import xlsxwriter as _xlsxwriter
                bio = BytesIO()
                workbook = _xlsxwriter.Workbook(bio)
                ws = workbook.add_worksheet()
                for c_idx, h in enumerate(headers):
                    ws.write(0, c_idx, h)
                for r_idx, r in enumerate(rows, start=1):
                    for c_idx, val in enumerate(r):
                        ws.write(r_idx, c_idx, val)
                workbook.close()
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'donations.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل حمایت‌ها", force_document=True)
                await event.edit("✅ فایل اکسل حمایت‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import csv
                with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, encoding='utf-8', newline='') as tmp:
                    writer = csv.writer(tmp)
                    writer.writerow(headers)
                    for r in rows:
                        writer.writerow(r)
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📥 فایل CSV حمایت‌ها (fallback)", force_document=True)
                await event.edit("✅ فایل CSV حمایت‌ها ارسال شد (fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return

        elif data == "admin_export_tickets_excel":
            await event.edit("⏳ در حال تهیه فایل اکسل تیکت‌ها...", buttons=CANCEL_BUTTON)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, message, status, created_at FROM tickets ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            headers = ["شماره", "کاربر", "پیام", "وضعیت", "تاریخ"]
            if PD_AVAILABLE:
                try:
                    df = safe_df_from_rows(rows, headers)
                    bio = BytesIO()
                    df.to_excel(bio, index=False, engine='openpyxl')
                    bio.seek(0)
                    bio = apply_persian_xlsx_style(bio)
                    try:
                        bio.name = 'tickets.xlsx'
                    except Exception:
                        pass
                    await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل تیکت‌ها", force_document=True)
                    await event.edit("✅ فایل اکسل تیکت‌ها ارسال شد.", buttons=get_admin_main_menu())
                    return
                except Exception:
                    pass
            try:
                from openpyxl import Workbook
                wb = Workbook()
                ws = wb.active
                ws.append(headers)
                for r in rows:
                    ws.append(list(r))
                bio = BytesIO()
                wb.save(bio)
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'tickets.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل تیکت‌ها", force_document=True)
                await event.edit("✅ فایل اکسل تیکت‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import xlsxwriter as _xlsxwriter
                bio = BytesIO()
                workbook = _xlsxwriter.Workbook(bio)
                ws = workbook.add_worksheet()
                for c_idx, h in enumerate(headers):
                    ws.write(0, c_idx, h)
                for r_idx, r in enumerate(rows, start=1):
                    for c_idx, val in enumerate(r):
                        ws.write(r_idx, c_idx, val)
                workbook.close()
                bio.seek(0)
                bio = apply_persian_xlsx_style(bio)
                try:
                    bio.name = 'tickets.xlsx'
                except Exception:
                    pass
                await client.send_file(event.chat_id, bio, caption="📥 فایل اکسل تیکت‌ها", force_document=True)
                await event.edit("✅ فایل اکسل تیکت‌ها ارسال شد.", buttons=get_admin_main_menu())
                return
            except Exception:
                pass
            try:
                import csv
                with tempfile.NamedTemporaryFile(suffix='.csv', mode='w', delete=False, encoding='utf-8', newline='') as tmp:
                    writer = csv.writer(tmp)
                    writer.writerow(headers)
                    for r in rows:
                        writer.writerow(r)
                    tmp.flush()
                    await client.send_file(event.chat_id, tmp.name, caption="📥 فایل CSV تیکت‌ها (fallback)", force_document=True)
                await event.edit("✅ فایل CSV تیکت‌ها ارسال شد (fallback).", buttons=get_admin_main_menu())
            except Exception as e:
                await event.edit(f"❌ خطا در تولید فایل: {e}", buttons=get_admin_main_menu())
            return

        elif data == "admin_message_members":
            buttons = [
                [Button.inline("✅ تاییدشدگان", b"admin_message_members_approved")],
                [Button.inline("❌ ردشدگان", b"admin_message_members_rejected")],
                [Button.inline("👥 همه اعضا", b"admin_message_members_all")],
                [Button.inline("🔙 بازگشت", b"admin_settings")],
            ]
            await event.edit("✉️ به کدام گروه از اعضا پیام ارسال شود؟", buttons=buttons)
            return
        elif data == "admin_send_to_id":
            # prompt admin to enter numeric Telegram user id
            await event.edit("✉️ لطفا آیدی عددی کاربر را وارد کنید (مثال: 123456789):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_target_id")
            return
        elif data == "admin_edit_user_help":
            help_text = get_setting('user_help_text', 'راهنمای جامع ربات هنوز توسط ادمین تنظیم نشده است.')
            await event.edit("✏️ لطفاً متن راهنمای جامع ربات را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_user_help_text", {"current": help_text})
            return
        elif data == "admin_edit_membership_desc":
            membership_desc = get_setting('membership_description', 'برای عضویت، لطفاً اطلاعات خود را کامل وارد کنید. پس از بررسی، نتیجه به شما اطلاع داده خواهد شد.')
            await event.edit("✏️ لطفاً متن توضیحات عضویت را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_membership_desc", {"current": membership_desc})
            return
            await event.edit("⚙️ تنظیمات اعلان‌ها:", buttons=buttons)

        elif data == "toggle_single_reg":
            if not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting('single_registration_per_user', '1')
            new = '0' if cur == '1' else '1'
            set_setting('single_registration_per_user', new)
            label = "✅ روشن" if new == '1' else '❌ خاموش'
            try:
                await event.answer(f"🔄 محدودیت یک ثبت‌نام/کاربر: {label}", alert=True)
            except Exception:
                pass
            # refresh admin settings view
            nreg = "✅ روشن" if get_setting("notify_new_registration","1") == "1" else "❌ خاموش"
            ntick = "✅ روشن" if get_setting("notify_new_ticket","1") == "1" else "❌ خاموش"
            single_reg = "✅ روشن" if get_setting("single_registration_per_user","1") == "1" else "❌ خاموش"
            try:
                donate_card = get_setting('donation_card_number', '—')
            except Exception:
                donate_card = '—'
            donate_label = donate_card if donate_card else '—'
            buttons = [
                [Button.inline(f"ثبت‌نام جدید: {nreg}", b"toggle_notify_reg")],
                [Button.inline(f"تیکت جدید: {ntick}", b"toggle_notify_ticket")],
                [Button.inline(f"🔒 یک ثبت‌نام/کاربر: {single_reg}", b"toggle_single_reg")],
                [Button.inline(f"💳 شماره کارت حمایت: {donate_label}", b"admin_set_donation_card")],
                [Button.inline("🧾 صاحب کارت (ویرایش)", b"admin_set_donation_holder")],
                [Button.inline("📝 توضیحات حمایت (ویرایش)", b"admin_set_donation_desc")],
                [Button.inline("🏠 بازگشت", b"main_menu")],
                [Button.inline("⭐ مدیریت میانبرهای رویداد", b"admin_manage_main_events")]
            ]
            await event.edit("⚙️ تنظیمات اعلان‌ها:", buttons=buttons)
        # toggle_console_logs handler removed

        elif data == "toggle_notify_reg":
            # allow OWNER or any admin to toggle notification setting
            if not (user_id == OWNER_ID or utils_is_admin(DB_NAME, user_id)):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting("notify_new_registration","1")
            set_setting("notify_new_registration", "0" if cur == "1" else "1")
            
            nreg = "✅ روشن" if get_setting("notify_new_registration","1") == "1" else "❌ خاموش"
            ntick = "✅ روشن" if get_setting("notify_new_ticket","1") == "1" else "❌ خاموش"
            buttons = [
                [Button.inline(f"ثبت‌نام جدید: {nreg}", b"toggle_notify_reg")],
                [Button.inline(f"تیکت جدید: {ntick}", b"toggle_notify_ticket")],
                [Button.inline("🏠 بازگشت", b"main_menu")],
                [Button.inline("⭐ مدیریت میانبرهای رویداد", b"admin_manage_main_events")]
            ]
            await event.edit("⚙️ تنظیمات اعلان‌ها:", buttons=buttons)

        elif data == "toggle_notify_membership":
            if not (user_id == OWNER_ID or utils_is_admin(DB_NAME, user_id)):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting("notify_new_membership","1")
            set_setting("notify_new_membership", "0" if cur == "1" else "1")
            # refresh
            nreg = "✅ روشن" if get_setting("notify_new_registration","1") == "1" else "❌ خاموش"
            ntick = "✅ روشن" if get_setting("notify_new_ticket","1") == "1" else "❌ خاموش"
            nmemb = "✅ روشن" if get_setting("notify_new_membership","1") == "1" else "❌ خاموش"
            nidea = "✅ روشن" if get_setting("notify_new_idea","1") == "1" else "❌ خاموش"
            ncollab = "✅ روشن" if get_setting("notify_new_collab","1") == "1" else "❌ خاموش"
            ndonate = "✅ روشن" if get_setting("notify_new_donation","1") == "1" else "❌ خاموش"
            single_reg = "✅ روشن" if get_setting("single_registration_per_user","1") == "1" else "❌ خاموش"
            buttons = [
                [Button.inline(f"ثبت‌نام جدید: {nreg}", b"toggle_notify_reg")],
                [Button.inline(f"تیکت جدید: {ntick}", b"toggle_notify_ticket")],
                [Button.inline(f"درخواست عضویت: {nmemb}", b"toggle_notify_membership")],
                [Button.inline(f"ایده جدید: {nidea}", b"toggle_notify_ideas")],
                [Button.inline(f"درخواست همکاری: {ncollab}", b"toggle_notify_collabs")],
                [Button.inline(f"حمایت جدید: {ndonate}", b"toggle_notify_donations")],
                [Button.inline(f"🔒 یک ثبت‌نام/کاربر: {single_reg}", b"toggle_single_reg")],
                [Button.inline("✏️ ویرایش", b"admin_edit_menu")],
                [Button.inline("🏠 بازگشت", b"main_menu")]
            ]
            await event.edit("⚙️ تنظیمات اعلان‌ها:", buttons=buttons)

        elif data == "toggle_notify_ideas":
            if not (user_id == OWNER_ID or utils_is_admin(DB_NAME, user_id)):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting("notify_new_idea","1")
            set_setting("notify_new_idea", "0" if cur == "1" else "1")
            await event.edit("تنظیم ذخیره شد.", buttons=get_admin_main_menu())

        elif data == "toggle_notify_collabs":
            if not (user_id == OWNER_ID or utils_is_admin(DB_NAME, user_id)):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting("notify_new_collab","1")
            set_setting("notify_new_collab", "0" if cur == "1" else "1")
            await event.edit("تنظیم ذخیره شد.", buttons=get_admin_main_menu())

        elif data == "toggle_notify_donations":
            if not (user_id == OWNER_ID or utils_is_admin(DB_NAME, user_id)):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting("notify_new_donation","1")
            set_setting("notify_new_donation", "0" if cur == "1" else "1")
            await event.edit("تنظیم ذخیره شد.", buttons=get_admin_main_menu())

        elif data == "toggle_notify_ticket":
            # allow OWNER or any admin to toggle notification setting
            if not (user_id == OWNER_ID or utils_is_admin(DB_NAME, user_id)):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            cur = get_setting("notify_new_ticket","1")
            set_setting("notify_new_ticket", "0" if cur == "1" else "1")
            
            nreg = "✅ روشن" if get_setting("notify_new_registration","1") == "1" else "❌ خاموش"
            ntick = "✅ روشن" if get_setting("notify_new_ticket","1") == "1" else "❌ خاموش"
            buttons = [
                [Button.inline(f"ثبت‌نام جدید: {nreg}", b"toggle_notify_reg")],
                [Button.inline(f"تیکت جدید: {ntick}", b"toggle_notify_ticket")],
                [Button.inline("🏠 بازگشت", b"main_menu")],
                [Button.inline("⭐ مدیریت میانبرهای رویداد", b"admin_manage_main_events")]
            ]
            await event.edit("⚙️ تنظیمات اعلان‌ها:", buttons=buttons)

        elif data == "admin_set_donation_card":
            if user_id != OWNER_ID and not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            await event.edit("💳 لطفا شماره کارت حمایت را ارسال کنید (فقط اعداد و فاصله مجاز است):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_donation_card")
            return

        elif data == "admin_set_donation_desc":
            if user_id != OWNER_ID and not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            await event.edit("📝 لطفا متن توضیحات حمایت را ارسال کنید (قابل نمایش در صفحه حمایت کاربران):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_donation_desc")
            return

        elif data == "admin_set_donation_holder":
            if user_id != OWNER_ID and not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            await event.edit("🧾 لطفا نام صاحب کارت را ارسال کنید (مثلا: نام و نام خانوادگی):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_donation_holder")
            return

        elif data == "admin_manage_main_events":
            
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events ORDER BY id DESC LIMIT 200")
            evs = c.fetchall()
            conn.close()
            
            raw = get_setting('main_events', '')
            selected = [int(x) for x in raw.split(',') if x.strip().isdigit()]
            buttons = []
            for eid, title in evs:
                prefix = "✅ " if eid in selected else "  "
                buttons.append([Button.inline(f"{prefix}{title}", f"admin_toggle_main_{eid}")])
            buttons.append([Button.inline("🔙 بازگشت", b"admin_settings")])
            await event.edit("⭐ مدیریت میانبرهای اصلی (تا 3 رویداد):\nبرای فعال/غیرفعال کردن روی عنوان کلیک کنید.", buttons=buttons)

        elif data.startswith("admin_toggle_main_") and not data.startswith("admin_toggle_main_edit_"):
            eid = int(data.split("_")[-1])
            raw = get_setting('main_events', '')
            selected = [int(x) for x in raw.split(',') if x.strip().isdigit()]
            if eid in selected:
                selected.remove(eid)
            else:
                if len(selected) >= 3:
                    await event.answer("❌ حداکثر 3 میانبر مجاز است.", alert=True)
                    return
                selected.append(eid)
            set_setting('main_events', ','.join(map(str, selected)))
            
            try:
                mtext = getattr(event.message, 'message', '') or ''
            except Exception:
                mtext = ''
            
            if '⭐ مدیریت میانبرهای اصلی' in mtext:
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("SELECT id, title FROM events ORDER BY id DESC LIMIT 200")
                evs = c.fetchall()
                conn.close()
                raw2 = get_setting('main_events', '')
                selected2 = [int(x) for x in raw2.split(',') if x.strip().isdigit()]
                buttons = []
                for _eid, title in evs:
                    prefix = "✅ " if _eid in selected2 else "  "
                    buttons.append([Button.inline(f"{prefix}{title}", f"admin_toggle_main_{_eid}")])
                buttons.append([Button.inline("🔙 بازگشت", b"admin_settings")])
                await event.edit("⭐ مدیریت میانبرهای اصلی (تا 3 رویداد):\nبرای فعال/غیرفعال کردن روی عنوان کلیک کنید.", buttons=buttons)
                return
            
            if 'لطفا بخش مورد نظر برای ویرایش' in mtext:
                
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("SELECT title, description, cost_type, card_number, is_active FROM events WHERE id = ?", (eid,))
                result = c.fetchone()
                conn.close()
                title = result[0] if result else f"رویداد #{eid}"
                is_active = result[4] if result and len(result) > 4 else 0
                status_text = "✅ فعال" if is_active else "❌ غیرفعال"
                raw3 = get_setting('main_events', '')
                selected3 = [int(x) for x in raw3.split(',') if x.strip().isdigit()]
                main_label = "⭐ میانبر (فعال)" if eid in selected3 else "⭐ میانبر (غیر فعال)"
                buttons = [
                    [Button.inline(main_label, f"admin_toggle_main_{eid}")],
                    [Button.inline("✏️ ویرایش عنوان", f"edit_title_{eid}"), Button.inline("📝 ویرایش توضیحات", f"edit_desc_{eid}")],
                    [Button.inline("💰 ویرایش هزینه", f"edit_cost_{eid}"), Button.inline("💳 ویرایش شماره کارت", f"edit_card_{eid}")],
                    [Button.inline("🖼️ ویرایش پوستر", f"edit_poster_{eid}"), Button.inline("🎯 تنظیم ظرفیت", f"admin_set_capacity_{eid}")],
                    [Button.inline(f"🔄 وضعیت ({status_text})", f"admin_toggle_event_{eid}"), Button.inline("⏱️ تنظیم مهلت (شمسی)", f"admin_set_deadline_{eid}")],
                    [Button.inline("📑 تنظیم گزارش کار", f"admin_set_report_{eid}"), Button.inline("🗂️ مدیریت گزارش‌ها", f"admin_manage_reports_{eid}")],
                    [Button.inline("🗑️ حذف کامل رویداد", f"admin_delete_event_{eid}")],
                    [Button.inline("✅ تایید گروهی در انتظار", f"admin_bulk_approve_{eid}"), Button.inline("🔔 یادآوری به تاییدشدگان", f"admin_remind_{eid}")],
                    [Button.inline("✉️ پیام به تاییدشدگان", f"admin_message_approved_{eid}"), Button.inline("✉️ پیام به ردشدگان", f"admin_message_rejected_{eid}")],
                    [Button.inline("🔙 بازگشت", b"admin_manage_events_0"), Button.inline("🏠 منو", b"main_menu")]
                ]
                await event.edit("لطفا بخش مورد نظر برای ویرایش را انتخاب کنید:", buttons=buttons)
                return
            
            await event.answer("✅ تغییر ذخیره شد.", alert=True)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events ORDER BY id DESC LIMIT 200")
            evs = c.fetchall()
            conn.close()
            raw2 = get_setting('main_events', '')
            selected2 = [int(x) for x in raw2.split(',') if x.strip().isdigit()]
            buttons = []
            for _eid, title in evs:
                prefix = "✅ " if _eid in selected2 else "  "
                buttons.append([Button.inline(f"{prefix}{title}", f"admin_toggle_main_{_eid}")])
            buttons.append([Button.inline("🔙 بازگشت", b"admin_settings")])
            await event.edit("⭐ مدیریت میانبرهای اصلی (تا 3 رویداد):\nبرای فعال/غیرفعال کردن روی عنوان کلیک کنید.", buttons=buttons)

        elif data.startswith("admin_toggle_main_edit_"):
            
            eid = int(data.split("_")[-1])
            raw = get_setting('main_events', '')
            selected = [int(x) for x in raw.split(',') if x.strip().isdigit()]
            if eid in selected:
                selected.remove(eid)
            else:
                if len(selected) >= 3:
                    await event.answer("❌ حداکثر 3 میانبر مجاز است.", alert=True)
                    return
                selected.append(eid)
            set_setting('main_events', ','.join(map(str, selected)))
            
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT title, description, cost_type, card_number, is_active, poster_file_id FROM events WHERE id = ?", (eid,))
            result = c.fetchone()
            conn.close()
            if not result:
                await event.answer("❌ رویداد یافت نشد!", alert=True)
                return
            title, desc, cost_type, card, is_active, poster_path = result
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT fixed_cost, student_cost, non_student_cost, capacity, end_at_ts FROM events WHERE id = ?", (eid,))
            extra = c.fetchone()
            conn.close()
            fixed_cost = extra[0] if extra else 0
            student_cost = extra[1] if extra else 0
            non_student_cost = extra[2] if extra else 0
            capacity = extra[3] if extra else None
            end_at_ts = extra[4] if extra else None
            cost_display = "رایگان"
            if cost_type == "fixed":
                cost_display = f"{fixed_cost:,} تومان"
            elif cost_type == "variable":
                cost_display = f"دانشجو: {student_cost:,} — غیر دانشجو: {non_student_cost:,} تومان"
            status_text = "✅ فعال" if is_active else "❌ غیرفعال"
            cap_txt = "بدون محدودیت" if (capacity is None or capacity == -1) else str(capacity)
            deadline_txt = "—"
            if end_at_ts:
                try:
                    from utils import epoch_to_jalali_str
                    deadline_txt = epoch_to_jalali_str(end_at_ts) + " (ساعت رسمی ایران)"
                except Exception:
                    deadline_txt = str(end_at_ts)
            preview = f"""
📌 {title}
{desc or '—'}

وضعیت: {status_text}
ظرفیت: {cap_txt}
هزینه: {cost_display}
کارت: {card or '—'}
مهلت: {deadline_txt}
""".strip()
            try:
                if poster_path and is_safe_upload_path(poster_path):
                    await client.send_file(event.chat_id, poster_path, caption=preview)
                else:
                    await client.send_message(event.chat_id, preview)
            except Exception:
                pass
            raw_main = get_setting('main_events', '')
            selected_main = [int(x) for x in raw_main.split(',') if x.strip().isdigit()]
            main_label = "⭐ میانبر (فعال)" if eid in selected_main else "⭐ میانبر (غیر فعال)"
            buttons = [
                [Button.inline(main_label, f"admin_toggle_main_edit_{eid}")],
                [Button.inline("✏️ ویرایش عنوان", f"edit_title_{eid}"), Button.inline("📝 ویرایش توضیحات", f"edit_desc_{eid}")],
                [Button.inline("💰 ویرایش هزینه", f"edit_cost_{eid}"), Button.inline("💳 ویرایش شماره کارت", f"edit_card_{eid}")],
                [Button.inline("🖼️ ویرایش پوستر", f"edit_poster_{eid}"), Button.inline("🎯 تنظیم ظرفیت", f"admin_set_capacity_{eid}")],
                [Button.inline(f"🔄 وضعیت ({status_text})", f"admin_toggle_event_{eid}"), Button.inline("⏱️ تنظیم مهلت (شمسی)", f"admin_set_deadline_{eid}")],
                [Button.inline("📑 تنظیم گزارش کار", f"admin_set_report_{eid}"), Button.inline("🗂️ مدیریت گزارش‌ها", f"admin_manage_reports_{eid}")],
                [Button.inline("🗑️ حذف کامل رویداد", f"admin_delete_event_{eid}")],
                [Button.inline("✅ تایید گروهی در انتظار", f"admin_bulk_approve_{eid}"), Button.inline("🔔 یادآوری به تاییدشدگان", f"admin_remind_{eid}")],
                [Button.inline("✉️ پیام به تاییدشدگان", f"admin_message_approved_{eid}"), Button.inline("✉️ پیام به ردشدگان", f"admin_message_rejected_{eid}")],
                [Button.inline("🔙 بازگشت", b"admin_manage_events_0"), Button.inline("🏠 منو", b"main_menu")]
            ]
            await event.edit("لطفا بخش مورد نظر برای ویرایش را انتخاب کنید:", buttons=buttons)

        elif data == "maint_clear_pending_regs":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("DELETE FROM registrations WHERE status = 'pending'")
            deleted = c.rowcount
            conn.commit()
            conn.close()
            await event.answer(f"✅ {deleted} ثبت‌نام در انتظار حذف شد.", alert=True)

        elif data.startswith("maint_purge_event_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT poster_file_id, report_payloads FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            poster_path = row[0] if row else None
            payloads = row[1] if row else None
            try:
                if poster_path and is_safe_upload_path(poster_path):
                    os.remove(poster_path)
            except Exception:
                pass
            if payloads:
                try:
                    import json as _json
                    for p in _json.loads(payloads):
                        if p.get("type") == "file":
                            pth = p.get("path")
                            if pth and is_safe_upload_path(pth):
                                try:
                                    os.remove(pth)
                                except Exception:
                                    pass
                except Exception:
                    pass
            c.execute("DELETE FROM certificates WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM registrations WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM resources WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM attendance WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM events WHERE id = ?", (event_id,))
            conn.commit()
            conn.close()
            await event.answer("🧹 داده‌های رویداد حذف شد.", alert=True)

        elif data.startswith("maint_purge_event_files_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT poster_file_id, report_payloads FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            poster_path = row[0] if row else None
            payloads = row[1] if row else None
            try:
                if poster_path and is_safe_upload_path(poster_path):
                    os.remove(poster_path)
            except Exception:
                pass
            if payloads:
                try:
                    import json as _json
                    for p in _json.loads(payloads):
                        if p.get("type") == "file":
                            pth = p.get("path")
                            if pth and is_safe_upload_path(pth):
                                try:
                                    os.remove(pth)
                                except Exception:
                                    pass
                except Exception:
                    pass
            c.execute("SELECT file_id FROM certificates WHERE event_id = ?", (event_id,))
            for (fid,) in c.fetchall():
                try:
                    if fid and is_safe_upload_path(fid):
                        os.remove(fid)
                except Exception:
                    pass
            c.execute("UPDATE events SET poster_file_id = NULL WHERE id = ?", (event_id,))
            c.execute("UPDATE events SET report_payloads = NULL WHERE id = ?", (event_id,))
            c.execute("UPDATE registrations SET payment_receipt_file_id = NULL WHERE event_id = ?", (event_id,))
            c.execute("UPDATE certificates SET file_id = NULL WHERE event_id = ?", (event_id,))
            conn.commit()
            conn.close()
            await event.answer("🧺 فایل‌های رویداد حذف شد و ارجاعات پاک شدند.", alert=True)

        elif data == "maint_clear_closed_tickets":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("DELETE FROM tickets WHERE status = 'closed'")
            deleted = c.rowcount
            conn.commit()
            conn.close()
            await event.answer(f"✅ {deleted} تیکت بسته حذف شد.", alert=True)

        elif data == "maint_clear_orphan_resources":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, file_path FROM resources")
            rows = c.fetchall()
            removed = 0
            for rid, path in rows:
                if not os.path.exists(path or ""):
                    c.execute("DELETE FROM resources WHERE id = ?", (rid,))
                    removed += 1
            conn.commit()
            conn.close()
            await event.answer(f"✅ {removed} منبع یتیم حذف شد.", alert=True)

        elif data == "maint_clear_orphan_receipts":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, payment_receipt_file_id FROM registrations WHERE payment_receipt_file_id IS NOT NULL")
            rows = c.fetchall()
            removed = 0
            for rid, path in rows:
                if path and not os.path.exists(path):
                    c.execute("DELETE FROM registrations WHERE id = ?", (rid,))
                    removed += 1
            conn.commit()
            conn.close()
            await event.answer(f"✅ {removed} رسید یتیم/ناقص حذف شد.", alert=True)

        elif data == "maint_clear_stray_files":
            base = os.getcwd()
            safe_dirs = {os.path.join(base, "uploads")}
            safe_exts = {".db", ".log"}
            protected_roots = {base}
            removed = 0
            errors = 0
            for root, dirs, files in os.walk(base):
                rel = os.path.relpath(root, base)
                if rel.startswith('.') or rel.startswith(".git"):
                    continue
                for f in files:
                    path = os.path.join(root, f)
                    if path.startswith(os.path.join(base, "uploads")):
                        try:
                            os.remove(path)
                            removed += 1
                        except Exception:
                            errors += 1
                        continue
                    if os.path.splitext(path)[1].lower() in safe_exts:
                        continue
                    if path.endswith('.py'):
                        continue
                    try:
                        os.remove(path)
                        removed += 1
                    except Exception:
                        errors += 1
            await event.answer(f"🧹 حذف فایل‌ها به پایان رسید. حذف: {removed} | خطا: {errors}", alert=True)

        elif data == "admin_capacity":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title, COALESCE(capacity, -1) FROM events ORDER BY id DESC")
            events_list = c.fetchall()
            conn.close()
            if not events_list:
                await event.edit("📭 هیچ رویدادی وجود ندارد.", buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])
                return
            buttons = []
            for eid, title, cap in events_list:
                cap_txt = "بدون محدودیت" if cap is None or cap == -1 else str(cap)
                buttons.append([Button.inline(f"{title} — ظرفیت: {cap_txt}", f"admin_set_capacity_{eid}")])
            buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit("🎯 انتخاب رویداد برای تنظیم ظرفیت:", buttons=buttons)

        elif data == "admin_new_event_step1":
            await event.edit("📌 لطفا عنوان رویداد را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_new_event_title")

        elif data == "admin_stats":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events ORDER BY id DESC")
            events_list = c.fetchall()
            conn.close()
            buttons = []
            for eid, title in events_list:
                buttons.append([Button.inline(f"📊 {title}", f"stats_event_{eid}")])
            buttons.append([Button.inline("🏠 بازگشت به منو", b"main_menu")])
            await event.edit("📈 آمار بر اساس رویداد:", buttons=buttons)

        elif data.startswith("stats_event_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("""
                SELECT 
                    SUM(CASE WHEN r.status='approved' THEN 1 ELSE 0 END) AS approved_cnt,
                    SUM(CASE WHEN r.status='rejected' THEN 1 ELSE 0 END) AS rejected_cnt,
                    SUM(CASE WHEN r.status='pending' THEN 1 ELSE 0 END) AS pending_cnt
                FROM registrations r
                WHERE r.event_id = ?
            """, (event_id,))
            row = c.fetchone()
            approved_cnt, rejected_cnt, pending_cnt = row if row else (0, 0, 0)
            c.execute("SELECT title FROM events WHERE id = ?", (event_id,))
            title_row = c.fetchone()
            conn.close()
            title = title_row[0] if title_row else str(event_id)
            msg = f"""
📊 آمار رویداد: {title}

✅ تایید شده: {approved_cnt or 0}
❌ رد شده: {rejected_cnt or 0}
⏳ در انتظار: {pending_cnt or 0}
            """.strip()
            await event.edit(msg, buttons=[
                [Button.inline("🔙 بازگشت", b"admin_stats")],
                [Button.inline("🏠 منوی اصلی", b"main_menu")]
            ])

        elif data.startswith("admin_manage_events_"):
            page = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title, is_active FROM events ORDER BY id DESC")
            events_list = c.fetchall()
            conn.close()
            if not events_list:
                await event.edit("📭 هیچ رویدادی وجود ندارد.", buttons=[[Button.inline("➕ ثبت جدید", b"admin_new_event_step1")], [Button.inline("🏠 منو", b"main_menu")]])
                return
            per_page = 5
            total = len(events_list)
            start = page * per_page
            end = start + per_page
            slice_events = events_list[start:end]
            buttons = []
            for eid, title, is_active in slice_events:
                status = "✅ فعال" if is_active else "❌ غیرفعال"
                buttons.append([Button.inline(f"{title} — {status}", f"admin_edit_event_{eid}")])
            nav = []
            if page > 0:
                nav.append(Button.inline("⬅️ قبلی", f"admin_manage_events_{page-1}"))
            if end < total:
                nav.append(Button.inline("➡️ بعدی", f"admin_manage_events_{page+1}"))
            if nav:
                buttons.append(nav)
            buttons.append([Button.inline("🏠 بازگشت به منو", b"main_menu")])
            await event.edit("⚙️ رویدادهای ثبت شده:", buttons=buttons)

        elif data.startswith("admin_toggle_single_reg_event_"):
            event_id = int(data.split("_")[-1])
            if not utils_is_admin(DB_NAME, user_id):
                await event.answer("❌ دسترسی ندارید!", alert=True)
                return
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT single_registration FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            current = 1 if row and row[0] else 0
            new = 0 if current == 1 else 1
            c.execute("UPDATE events SET single_registration = ? WHERE id = ?", (new, event_id))
            conn.commit()
            conn.close()
            await event.answer("🔄 تنظیمات بروزرسانی شد.", alert=True)
            try:
                # refresh the edit view for the event
                await admin_callback_handler(event)
            except Exception:
                pass

        elif data.startswith("admin_edit_event_"):
            event_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT title, description, cost_type, card_number, is_active, poster_file_id FROM events WHERE id = ?", (event_id,))
            result = c.fetchone()
            conn.close()
            if not result:
                await event.answer("❌ رویداد یافت نشد!", alert=True)
                return
            title, desc, cost_type, card, is_active, poster_path = result
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT fixed_cost, student_cost, non_student_cost, capacity, end_at_ts FROM events WHERE id = ?", (event_id,))
            extra = c.fetchone()
            conn.close()
            fixed_cost = extra[0] if extra else 0
            student_cost = extra[1] if extra else 0
            non_student_cost = extra[2] if extra else 0
            capacity = extra[3] if extra else None
            end_at_ts = extra[4] if extra else None
            cost_display = "رایگان"
            if cost_type == "fixed":
                cost_display = f"{fixed_cost:,} تومان"
            elif cost_type == "variable":
                cost_display = f"دانشجو: {student_cost:,} — غیر دانشجو: {non_student_cost:,} تومان"
            status_text = "✅ فعال" if is_active else "❌ غیرفعال"
            cap_txt = "بدون محدودیت" if (capacity is None or capacity == -1) else str(capacity)
            deadline_txt = "—"
            if end_at_ts:
                try:
                    import datetime
                    deadline_txt = datetime.datetime.utcfromtimestamp(int(end_at_ts)).strftime("%Y-%m-%d %H:%M UTC")
                except Exception:
                    deadline_txt = str(end_at_ts)
            preview = f"""
📌 {title}
{desc or '—'}

وضعیت: {status_text}
ظرفیت: {cap_txt}
هزینه: {cost_display}
کارت: {card or '—'}
مهلت: {deadline_txt}
""".strip()
            try:
                if poster_path and is_safe_upload_path(poster_path):
                    await client.send_file(event.chat_id, poster_path, caption=preview)
                else:
                    await client.send_message(event.chat_id, preview)
            except Exception:
                pass
            status_text = "✅ فعال" if is_active else "❌ غیرفعال"
            try:
                raw_main = get_setting('main_events', '')
                selected_main = [int(x) for x in raw_main.split(',') if x.strip().isdigit()]
            except Exception:
                selected_main = []
            main_label = "⭐ میانبر (فعال)" if event_id in selected_main else "⭐ افزودن به میانبرها"
            try:
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("SELECT reminders_enabled FROM events WHERE id = ?", (event_id,))
                rrow = c.fetchone()
                conn.close()
                reminders_enabled = bool(rrow[0]) if rrow and rrow[0] else False
            except Exception:
                reminders_enabled = False
            # read per-event single-registration flag
            try:
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("SELECT single_registration FROM events WHERE id = ?", (event_id,))
                sr_row = c.fetchone()
                conn.close()
                single_registration = True if sr_row and sr_row[0] else False
            except Exception:
                single_registration = True

            single_label = "🔒 یک ثبت‌نام/کاربر (رویداد): ✅" if single_registration else "🔓 یک ثبت‌نام/کاربر (رویداد): ❌"

            buttons = [
                [Button.inline(main_label, f"admin_toggle_main_edit_{event_id}")],
                [Button.inline("✏️ ویرایش عنوان", f"edit_title_{event_id}"), Button.inline("📝 ویرایش توضیحات", f"edit_desc_{event_id}")],
                [Button.inline("💰 ویرایش هزینه", f"edit_cost_{event_id}"), Button.inline("💳 ویرایش شماره کارت", f"edit_card_{event_id}")],
                [Button.inline("🖼️ ویرایش پوستر", f"edit_poster_{event_id}"), Button.inline("🎯 تنظیم ظرفیت", f"admin_set_capacity_{event_id}")],
                [Button.inline(f"🔄 وضعیت ({status_text})", f"admin_toggle_event_{event_id}"), Button.inline("⏱️ تنظیم مهلت (شمسی)", f"admin_set_deadline_{event_id}")],
                [Button.inline("📑 تنظیم گزارش کار", f"admin_set_report_{event_id}"), Button.inline("🗂️ مدیریت گزارش‌ها", f"admin_manage_reports_{event_id}")],
                [Button.inline("🗂 مدیریت ثبت‌نام‌ها", f"admin_manage_event_regs_{event_id}")],
                [Button.inline(single_label, f"admin_toggle_single_reg_event_{event_id}")],
                [Button.inline("🗑️ حذف کامل رویداد", f"admin_delete_event_{event_id}"), Button.inline("🗑️ حذف مهلت", f"admin_confirm_clear_deadline_{event_id}")],
                [Button.inline("✅ تایید گروهی در انتظار", f"admin_bulk_approve_{event_id}"), Button.inline("🔔 یادآوری به تاییدشدگان", f"admin_remind_{event_id}")],
                [Button.inline("✉️ پیام به تاییدشدگان", f"admin_message_approved_{event_id}"), Button.inline("✉️ پیام به ردشدگان", f"admin_message_rejected_{event_id}" )],
                [Button.inline("🔙 بازگشت", b"admin_manage_events_0"), Button.inline("🏠 منو", b"main_menu")]
            ]
            await event.edit("لطفا بخش مورد نظر برای ویرایش را انتخاب کنید:", buttons=buttons)
        elif data.startswith("admin_delete_event_"):
            event_id = int(data.split("_")[-1])
            buttons = [
                [Button.inline("❌ لغو", b"main_menu"), Button.inline("🗑️ تأیید حذف", f"confirm_delete_event_{event_id}")]
            ]
            await event.edit("⚠️ آیا مطمئن هستید که می‌خواهید این رویداد و همه داده‌های مرتبط را حذف کنید؟ این عمل غیرقابل بازگشت است.", buttons=buttons)

        elif data.startswith("confirm_delete_event_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT poster_file_id, report_payloads FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            poster_path = row[0] if row else None
            payloads = row[1] if row else None
            try:
                if poster_path and is_safe_upload_path(poster_path):
                    os.remove(poster_path)
            except Exception:
                pass
            if payloads:
                try:
                    import json as _json
                    for p in _json.loads(payloads):
                        if p.get("type") == "file":
                            pth = p.get("path")
                            if pth and is_safe_upload_path(pth):
                                try:
                                    os.remove(pth)
                                except Exception:
                                    pass
                except Exception:
                    pass
            c.execute("DELETE FROM certificates WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM registrations WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM resources WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM attendance WHERE event_id = ?", (event_id,))
            c.execute("DELETE FROM events WHERE id = ?", (event_id,))
            conn.commit()
            conn.close()
            try:
                await event.answer("🗑️ رویداد و همه داده‌های مرتبط حذف شد.", alert=True)
            except:
                pass
            await event.edit("⚙️ رویدادهای ثبت شده:", buttons=[[Button.inline("🔙 بازگشت", b"admin_manage_events_0")],[Button.inline("🏠 منوی اصلی", b"main_menu")]])
        elif data.startswith("admin_confirm_clear_deadline_"):
            event_id = int(data.split("_")[-1])
            buttons = [[Button.inline("❌ لغو", b"main_menu"), Button.inline("🗑️ تأیید حذف مهلت", f"admin_clear_deadline_{event_id}")]]
            await event.edit("⚠️ آیا مطمئن هستید که می‌خواهید مهلت این رویداد را پاک کنید؟ این عمل قابل برگشت نیست.", buttons=buttons)

        elif data.startswith("admin_clear_deadline_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE events SET end_at_ts = NULL, end_set_by = NULL WHERE id = ?", (event_id,))
            conn.commit()
            conn.close()
            try:
                await event.answer("✅ مهلت رویداد پاک شد.", alert=True)
            except:
                pass
            try:
                await event.edit("مهلت حذف شد.", buttons=[[Button.inline("🔙 بازگشت", f"admin_edit_event_{event_id}")]])
            except Exception:
                pass
        elif data.startswith("edit_poster_"):
            event_id = int(data.split("_")[2])
            await event.edit("🖼️ لطفاً پوستر جدید را ارسال کنید (عکس)", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_new_poster", {"event_id": event_id})
        elif data.startswith("admin_set_deadline_"):
            event_id = int(data.split("_")[-1])
            await event.edit("⏱️ لطفا مهلت را به فرمت شمسی و به وقت رسمی ایران (مثال: 1403/07/01 18:30) ارسال کنید. ارقام فارسی پذیرفته می‌شوند.", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_deadline", {"event_id": event_id})

        elif data.startswith("edit_title_"):
            event_id = int(data.split("_")[2])
            await event.edit("📌 لطفا عنوان جدید را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "edit_event_title", {"event_id": event_id})

        elif data.startswith("edit_desc_"):
            event_id = int(data.split("_")[2])
            await event.edit("📝 لطفا توضیحات جدید را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "edit_event_desc", {"event_id": event_id})

        elif data.startswith("edit_cost_"):
            event_id = int(data.split("_")[2])
            buttons = [
                [Button.inline("رایگان", b"edit_cost_free")],
                [Button.inline("هزینه ثابت", b"edit_cost_fixed")],
                [Button.inline("هزینه متغیر", b"edit_cost_variable")],
                [Button.inline("🔙 بازگشت", f"admin_edit_event_{event_id}")]
            ]
            await event.edit("💰 نوع هزینه جدید را انتخاب کنید:", buttons=buttons)
            set_user_state(user_states, user_id, "edit_event_cost_type", {"event_id": event_id})

        elif data.startswith("edit_card_"):
            event_id = int(data.split("_")[2])
            await event.edit("💳 لطفا شماره کارت جدید را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "edit_event_card", {"event_id": event_id})

        elif data.startswith("admin_toggle_event_"):
            event_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE events SET is_active = 1 - is_active WHERE id = ?", (event_id,))
            conn.commit()
            conn.close()
            try:
                await event.answer("✅ وضعیت رویداد تغییر کرد.", alert=True)
            except:
                pass


        elif data.startswith("admin_set_report_"):
            event_id = int(data.split("_")[3])
            await event.edit("لطفا پیام‌های گزارش کار را از کانال فوروارد کنید (یک یا چند پیام). پس از اتمام، /done را ارسال کنید.", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_reports", {"event_id": event_id, "message_ids": [], "payloads": []})

        elif data.startswith("admin_manage_reports_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT report_message_ids, report_payloads, title FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            conn.close()
            msg_ids = row[0] if row else None
            payloads_json = row[1] if row else None
            title = row[2] if row and len(row) > 2 else f"رویداد #{event_id}"
            buttons = []
            if msg_ids or payloads_json:
                buttons.append([Button.inline("🧾 مشاهده آیتم‌های گزارش", f"admin_view_reports_{event_id}")])
                buttons.append([Button.inline("🗑️ پاک کردن گزارش‌ها", f"admin_clear_reports_{event_id}")])
            else:
                buttons.append([Button.inline("🏠 منو", b"main_menu")])
            summary = f"📌 مدیریت گزارش‌ها — {title}\n\n"
            if msg_ids:
                summary += f"🔢 آیدی پیام‌ها (نسخه پشتیبان): {msg_ids}\n"
            if payloads_json:
                try:
                    pl = json.loads(payloads_json)
                    summary += f"📦 تعداد آیتم‌ها: {len(pl)}\n"
                except Exception:
                    summary += "📦 آیتم‌ها: نامشخص\n"
            if not msg_ids and not payloads_json:
                summary += "📭 گزارشی تنظیم نشده است."
            await event.edit(summary, buttons=buttons)

        elif data.startswith("admin_view_reports_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT report_payloads FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            conn.close()
            payloads_json = row[0] if row else None
            if not payloads_json:
                await event.answer("📭 گزارشی برای این رویداد وجود ندارد.", alert=True)
                return
            try:
                payloads = json.loads(payloads_json)
            except Exception:
                await event.answer("❌ خطا در خواندن آیتم‌های گزارش.", alert=True)
                return
            for idx, p in enumerate(payloads):
                ptype = p.get('type')
                if ptype == 'forward':
                    ch = p.get('channel_id')
                    mid = p.get('message_id')
                    cap = p.get('caption')
                    text = f"#{idx+1} — فوروارد از کانال: channel_id={ch}, message_id={mid}"
                    if cap:
                        text += f"\nکپشن: {cap}"
                    try:
                        ent = await client.get_entity(int(ch))
                        uname = getattr(ent, 'username', None)
                        if uname:
                            text += f"\nلینک: https://t.me/{uname}/{int(mid)}"
                    except Exception:
                        pass
                    try:
                        await client.send_message(user_id, text)
                    except Exception:
                        pass
                elif ptype == 'file':
                    path = p.get('path')
                    cap = p.get('caption')
                    text = f"#{idx+1} — فایل محلی: {os.path.basename(path) if path else 'نامشخص'}"
                    if cap:
                        text += f"\nکپشن: {cap}"
                    try:
                        await client.send_message(user_id, text)
                    except Exception:
                        pass
                    try:
                        if path and is_safe_upload_path(path) and os.path.exists(path):
                            await client.send_file(user_id, path, caption=f"📁 فایل گزارش #{idx+1}: {os.path.basename(path)}")
                    except Exception:
                        pass
                elif ptype == 'text':
                    text = f"#{idx+1} — متن گزارش:\n{p.get('text','')}"
                    try:
                        await client.send_message(user_id, text)
                    except Exception:
                        pass
            await event.answer("✅ آیتم‌ها به ادمین ارسال شد.", alert=True)

        elif data.startswith("admin_clear_reports_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT report_payloads FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            payloads_json = row[0] if row else None
            if payloads_json:
                try:
                    pls = json.loads(payloads_json)
                    for p in pls:
                        if p.get('type') == 'file':
                            pth = p.get('path')
                            if pth and is_safe_upload_path(pth) and os.path.exists(pth):
                                try:
                                    os.remove(pth)
                                except Exception:
                                    pass
                except Exception:
                    pass
            c.execute("UPDATE events SET report_message_ids = NULL, report_payloads = NULL WHERE id = ?", (event_id,))
            conn.commit()
            conn.close()
            try:
                await event.answer("🧹 گزارش‌ها پاک شد.", alert=True)
            except Exception:
                pass
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT report_message_ids, report_payloads, title FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            conn.close()
            msg_ids = row[0] if row else None
            payloads_json = row[1] if row else None
            title = row[2] if row and len(row) > 2 else f"رویداد #{event_id}"
            buttons = []
            if msg_ids or payloads_json:
                buttons.append([Button.inline("🧾 مشاهده آیتم‌های گزارش", f"admin_view_reports_{event_id}")])
                buttons.append([Button.inline("🗑️ پاک کردن گزارش‌ها", f"admin_clear_reports_{event_id}")])
            else:
                buttons.append([Button.inline("🏠 منو", b"main_menu")])
            summary = f"📌 مدیریت گزارش‌ها — {title}\n\n"
            if msg_ids:
                summary += f"🔢 آیدی پیام‌ها (نسخه پشتیبان): {msg_ids}\n"
            if payloads_json:
                try:
                    pl = json.loads(payloads_json)
                    summary += f"📦 تعداد آیتم‌ها: {len(pl)}\n"
                except Exception:
                    summary += "📦 آیتم‌ها: نامشخص\n"
            if not msg_ids and not payloads_json:
                summary += "📭 گزارشی تنظیم نشده است."
            await event.edit(summary, buttons=buttons)

        elif data == "admin_pending_regs":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("""
                SELECT r.id, u.full_name, u.student_id, u.is_student, r.payment_receipt_file_id, e.title
                FROM registrations r
                JOIN users u ON r.user_id = u.user_id
                JOIN events e ON r.event_id = e.id
                WHERE r.status = 'pending'
            """)
            regs = c.fetchall()
            conn.close()
            if not regs:
                await event.edit("📭 هیچ ثبت‌نام در انتظاری وجود ندارد.", buttons=[[Button.inline("🏠 منو", b"main_menu")]])
                return
            buttons = []
            for reg_id, name, std_id, is_std, receipt, event_title in regs:
                student_type = "دانشجو" if is_std else "غیر دانشجو"
                buttons.append([Button.inline(f"{name} — {event_title}", f"admin_view_reg_{reg_id}")])
            buttons.append([Button.inline("🏠 بازگشت به منو", b"main_menu")])
            await event.edit("⏳ ثبت‌نام‌های در انتظار:", buttons=buttons)

        elif data.startswith("admin_set_capacity_"):
            event_id = int(data.split("_")[-1])
            await event.edit("🎯 لطفا ظرفیت رویداد را وارد کنید (عدد یا -1 برای بدون محدودیت):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_capacity", {"event_id": event_id})

        elif data.startswith("admin_bulk_approve_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE registrations SET status='approved' WHERE event_id = ? AND status='pending'", (event_id,))
            updated = c.rowcount
            c.execute("UPDATE users SET status='approved' WHERE user_id IN (SELECT user_id FROM registrations WHERE event_id = ?)", (event_id,))
            conn.commit()
            conn.close()
            await event.answer(f"✅ {updated} ثبت‌نام تایید شد.", alert=True)

        elif data.startswith("admin_remind_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("""
                SELECT DISTINCT r.user_id
                FROM registrations r
                WHERE LOWER(TRIM(r.status)) = 'approved' AND r.event_id = ?
            """, (event_id,))
            recipients = [row[0] for row in c.fetchall()]
            title_row = c.execute("SELECT title FROM events WHERE id = ?", (event_id,)).fetchone()
            title = title_row[0] if title_row else "رویداد"
            sample = []
            if recipients:
                placeholders = ','.join(['?'] * len(recipients))
                try:
                    q = f"SELECT user_id, full_name FROM users WHERE user_id IN ({placeholders}) LIMIT 6"
                    rows = c.execute(q, tuple(recipients)).fetchall()
                    sample = [f"{r[1] or 'کاربر'} ({r[0]})" for r in rows]
                except Exception:
                    sample = []
            conn.close()
            if not recipients:
                await event.answer("📭 کاربری برای یادآوری وجود ندارد.", alert=True)
                return
            sample_text = '\n'.join(sample) if sample else '—'
            body = f"⚠️ این پیام یادآوری برای رویداد '{title}' به {len(recipients)} کاربر ارسال خواهد شد.\n\nنمونه دریافت‌کنندگان:\n{sample_text}\n\nآیا ادامه می‌دهید؟"
            buttons = [[Button.inline("❌ لغو", b"main_menu"), Button.inline("🟢 ارسال یادآوری به تاییدشدگان", f"admin_remind_confirm_{event_id}")]]
            await event.edit(body, buttons=buttons)
            return

        elif data.startswith("admin_remind_confirm_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("""
                SELECT DISTINCT r.user_id
                FROM registrations r
                WHERE LOWER(TRIM(r.status)) = 'approved' AND r.event_id = ?
            """, (event_id,))
            recipients = [row[0] for row in c.fetchall()]
            title_row = c.execute("SELECT title FROM events WHERE id = ?", (event_id,)).fetchone()
            title = title_row[0] if title_row else "رویداد"
            conn.close()
            if not recipients:
                await event.answer("📭 کاربری برای یادآوری وجود ندارد.", alert=True)
                return
            msg = f"🔔 یادآوری: رویداد '{title}' نزدیک است. لطفاً اطلاعیه‌های کانال را دنبال کنید."
            sent = 0
            failed = 0
            from telethon import Button as TButton
            try:
                aconn = sqlite3.connect(DB_NAME)
                ac = aconn.cursor()
                ac.execute('''CREATE TABLE IF NOT EXISTS admin_actions (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    admin_id INTEGER NOT NULL,
                    action TEXT NOT NULL,
                    target_event INTEGER,
                    created_at TEXT DEFAULT (datetime('now','localtime'))
                )''')
                ac.execute('INSERT INTO admin_actions (admin_id, action, target_event) VALUES (?, ?, ?)', (user_id, 'send_reminder_to_approved', event_id))
                aconn.commit()
                aconn.close()
            except Exception:
                pass

            from utils import send_with_rate_limit
            import asyncio
            for uid in recipients:
                try:
                    btn = TButton.inline("مشاهده رویداد", f"event_{event_id}")
                    res = await send_with_rate_limit(client, uid, text=msg, buttons=[btn], delay_between=0.18)
                    if res:
                        sent += 1
                    else:
                        failed += 1
                        try:
                            lconn = sqlite3.connect(DB_NAME)
                            lc = lconn.cursor()
                            lc.execute('''CREATE TABLE IF NOT EXISTS send_errors (
                                id INTEGER PRIMARY KEY AUTOINCREMENT,
                                event_id INTEGER,
                                user_id INTEGER,
                                error TEXT,
                                created_at TEXT DEFAULT (datetime('now','localtime'))
                            )''')
                            lc.execute('INSERT INTO send_errors (event_id, user_id, error) VALUES (?, ?, ?)', (event_id, uid, 'send_failed'))
                            lconn.commit()
                            lconn.close()
                        except Exception:
                            pass
                except Exception:
                    failed += 1
                await asyncio.sleep(0.12)
            await event.answer(f"🔔 یادآوری برای {sent} نفر ارسال شد. ({failed} خطا)", alert=True)

        elif data.startswith("admin_message_approved_"):
            event_id = int(data.split("_")[-1])
            await event.edit("لطفا پیام، عکس یا ویدیوی خود را ارسال کنید (در انتهای پیام #پین برای پین کردن اضافه کنید):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_broadcast_content", {"target": "approved_event", "event_id": event_id})

        elif data.startswith("admin_message_rejected_"):
            event_id = int(data.split("_")[-1])
            await event.edit("لطفا پیام، عکس یا ویدیوی خود را ارسال کنید (در انتهای پیام #پین برای پین کردن اضافه کنید):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_broadcast_content", {"target": "rejected_event", "event_id": event_id})

        elif data.startswith("admin_view_reg_"):
            reg_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("""
                SELECT r.user_id, u.full_name, u.national_id, u.phone, u.is_student, u.student_id,
                       r.payment_receipt_file_id, e.title, e.cost_type, e.fixed_cost, e.student_cost, e.non_student_cost
                FROM registrations r
                JOIN users u ON r.user_id = u.user_id
                JOIN events e ON r.event_id = e.id
                WHERE r.id = ?
            """, (reg_id,))
            result = c.fetchone()
            conn.close()
            if not result:
                await event.answer("❌ ثبت‌نام یافت نشد!", alert=True)
                return
            (user_id, name, nid, phone, is_std, std_id, receipt, event_title, cost_type, fixed, std_cost, non_std_cost) = result
            cost_display = "رایگان"
            if cost_type == "fixed":
                cost_display = f"{fixed:,} تومان"
            elif cost_type == "variable":
                cost_display = f"{std_cost:,} تومان (دانشجو)" if is_std else f"{non_std_cost:,} تومان (غیر دانشجو)"
            msg = f"""
📌 ثبت‌نام: #{reg_id}
👤 نام: {name}
🆔 کد ملی: {nid}
📞 تلفن: {phone}
🎓 {('دانشجو — ' + std_id) if is_std else 'غیر دانشجو'}
📅 رویداد: {event_title}
💰 هزینه: {cost_display}
            """.strip()
            buttons = [
                [Button.inline("✅ تایید", f"admin_approve_reg_{reg_id}")],
                [Button.inline("❌ رد", f"admin_reject_reg_{reg_id}")],
                [Button.inline("🔙 بازگشت", b"admin_pending_regs")],
                [Button.inline("🏠 منو", b"main_menu")]
            ]
            if receipt and os.path.exists(receipt) and receipt.startswith("uploads" + os.sep):
                try:
                    await client.send_file(event.chat_id, receipt, caption="💳 رسید پرداخت:")
                except:
                    msg += "\n\n⚠️ رسید: [ارسال نشد — خطای فایل]"
            await event.edit(msg, buttons=buttons)

        elif data.startswith("admin_approve_reg_"):
            reg_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE registrations SET status = 'approved' WHERE id = ?", (reg_id,))
            c.execute("SELECT user_id FROM registrations WHERE id = ?", (reg_id,))
            user_id = c.fetchone()[0]
            c.execute("INSERT OR IGNORE INTO users (user_id, status) VALUES (?, 'approved')", (user_id,))
            c.execute("UPDATE users SET status = 'approved' WHERE user_id = ?", (user_id,))
            conn.commit()
            conn.close()
            try:
                await client.send_message(user_id, "✅ ثبت‌نام شما تایید شد! می‌توانید از بخش 'گواهی‌های من' پس از برگزاری رویداد، گواهی خود را دریافت کنید.")
            except:
                pass
            try:
                await event.answer("✅ ثبت‌نام تایید شد.", alert=True)
            except:
                pass

        elif data.startswith("admin_reject_reg_"):
            reg_id = int(data.split("_")[3])
            await event.edit("لطفا دلیل رد ثبت‌نام را بنویسید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_reject_reason", {"reg_id": reg_id})

        elif data == "admin_manage_admins":
            if user_id != OWNER_ID:
                await event.answer("❌ فقط ادمین اصلی می‌تواند مدیریت کند!", alert=True)
                return
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT a.user_id, u.full_name FROM admins a LEFT JOIN users u ON a.user_id = u.user_id")
            admins = c.fetchall()
            conn.close()
            buttons = [[Button.inline(f"➕ افزودن ادمین جدید", b"admin_add_admin")]]
            for admin_id, name in admins:
                if admin_id == OWNER_ID:
                    buttons.append([Button.inline(f"👑 {name or admin_id} (اصلی)", b"dummy")])
                else:
                    buttons.append([Button.inline(f"🗑️ حذف {name or admin_id}", f"admin_remove_admin_{admin_id}")])
            buttons.append([Button.inline("🏠 بازگشت به منو", b"main_menu")])
            await event.edit("👑 لیست ادمین‌ها:", buttons=buttons)

        # admin_actions_history handler removed

        elif data == "admin_membership_requests":
            # show pending membership requests (first page)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, full_name, status FROM memberships ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            if not rows:
                # No membership requests right now — still offer the Excel export button
                buttons = []
                buttons.append([Button.inline("📥 دریافت اکسل عضویت", b"admin_export_members_excel")])
                buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
                await event.edit("📭 هیچ درخواستی وجود ندارد.", buttons=buttons)
                return
            buttons = []
            for mid, name, status in rows[:20]:
                label = f"#{mid} — {name} — {status}"
                buttons.append([Button.inline(label, f"view_membership_{mid}")])
            # Add export excel and back buttons for membership section
            buttons.append([Button.inline("📥 دریافت اکسل عضویت", b"admin_export_members_excel")])
            buttons.append([Button.inline("✉️ پیام به اعضا", b"admin_message_members")])
            buttons.append([Button.inline("🔙 بازگشت", b"admin_settings"), Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit("📬 درخواست‌های عضویت:", buttons=buttons)

        elif data.startswith("view_membership_"):
            mid = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, full_name, major, entry_year, student_number, national_id, phone, telegram_username, student_card_file, status FROM memberships WHERE id = ?", (mid,))
            row = c.fetchone()
            conn.close()
            if not row:
                await event.answer("❌ درخواست یافت نشد!", alert=True)
                return
            user_id_t, full_name, major, entry_year, student_number, national_id, phone, tusername, card_file, status = row
            msg = f"📌 درخواست عضویت #{mid}\n👤 نام: {full_name}\n🎓 رشته: {major}\n📆 ورود: {entry_year}\n🆔 دانشجویی: {student_number}\n🔢 کدملی: {national_id}\n📞 تلفن: {phone}\n📨 آیدی: {tusername}\n\nوضعیت: {status}"
            buttons = [
                [Button.inline("✅ تایید", f"approve_membership_{mid}"), Button.inline("❌ رد", f"reject_membership_{mid}")],
                [Button.inline("🔙 بازگشت", b"admin_membership_requests"), Button.inline("🏠 منوی اصلی", b"main_menu")]
            ]
            try:
                if card_file and is_safe_upload_path(card_file) and os.path.exists(card_file):
                    await client.send_file(event.chat_id, card_file, caption=msg, buttons=buttons)
                else:
                    await event.edit(msg, buttons=buttons)
            except Exception:
                await event.edit(msg, buttons=buttons)

        elif data.startswith("approve_membership_"):
            mid = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id FROM memberships WHERE id = ?", (mid,))
            row = c.fetchone()
            if not row:
                conn.close()
                await event.answer("❌ درخواست یافت نشد!", alert=True)
                return
            uid = row[0]
            c.execute("UPDATE memberships SET status = 'approved' WHERE id = ?", (mid,))
            c.execute("INSERT OR IGNORE INTO users (user_id, status) VALUES (?, 'approved')", (uid,))
            c.execute("UPDATE users SET status = 'approved' WHERE user_id = ?", (uid,))
            conn.commit()
            conn.close()
            try:
                await client.send_message(uid, "✅ درخواست عضویت شما تایید شد!")
            except Exception:
                pass
            await event.answer("✅ درخواست تایید شد.", alert=True)
            await event.edit("✅ درخواست تایید شد.", buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])

        elif data.startswith("reject_membership_"):
            mid = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id FROM memberships WHERE id = ?", (mid,))
            row = c.fetchone()
            if not row:
                conn.close()
                await event.answer("❌ درخواست یافت نشد!", alert=True)
                return
            uid = row[0]
            c.execute("UPDATE memberships SET status = 'rejected' WHERE id = ?", (mid,))
            c.execute("UPDATE users SET status = 'rejected' WHERE user_id = ?", (uid,))
            conn.commit()
            conn.close()
            try:
                await client.send_message(uid, "❌ متأسفانه درخواست عضویت شما رد شد.")
            except Exception:
                pass
            await event.answer("✅ درخواست رد شد.", alert=True)
            await event.edit("✅ درخواست رد شد.", buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])

        elif data == "admin_add_admin":
            if user_id != OWNER_ID:
                await event.answer("❌ فقط ادمین اصلی می‌تواند اضافه کند!", alert=True)
                return
            await event.edit("لطفا آیدی عددی کاربر را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_user_id_to_add")

        elif data.startswith("admin_remove_admin_"):
            if user_id != OWNER_ID:
                await event.answer("❌ فقط ادمین اصلی می‌تواند حذف کند!", alert=True)
                return
            target_id = int(data.split("_")[3])
            if target_id == OWNER_ID:
                await event.answer("❌ نمی‌توان ادمین اصلی را حذف کرد!", alert=True)
                return
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("DELETE FROM admins WHERE user_id = ?", (target_id,))
            conn.commit()
            conn.close()
            await event.answer("✅ ادمین حذف شد.", alert=True)

        elif data == "admin_broadcast":
            buttons = [
                [Button.inline("📬 همه کاربران", b"broadcast_all"), Button.inline("✅ کاربران تایید شده", b"broadcast_approved")],
                [Button.inline("❌ کاربران رد شده", b"broadcast_rejected"), Button.inline("✉️ پیام به آیدی (خاص)", b"admin_send_to_id")],
                [Button.inline("🏠 بازگشت به منو", b"main_menu")]
            ]
            await event.edit("لطفا گروه هدف را انتخاب کنید:", buttons=buttons)

        elif data == "broadcast_by_event":
            await event.answer("این گزینه حذف شده است.", alert=True)

        elif data.startswith("broadcast_"):
            target = data.split("_")[1]
            if target == "approved":
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("SELECT id, title FROM events ORDER BY id DESC")
                events_list = c.fetchall()
                conn.close()
                if not events_list:
                    await event.edit("📭 هیچ رویدادی موجود نیست.", buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])
                    return
                buttons = []
                for eid, title in events_list:
                    buttons.append([Button.inline(title, f"broadcast_event_{eid}")])
                buttons.append([Button.inline("🏠 بازگشت", b"main_menu")])
                await event.edit("لطفا رویداد را برای ارسال به تاییدشدگان انتخاب کنید:", buttons=buttons)
            else:
                await event.edit("لطفا پیام، عکس یا ویدیوی خود را ارسال کنید (در انتهای پیام #پین برای پین کردن اضافه کنید):", buttons=CANCEL_BUTTON)
                set_user_state(user_states, user_id, "admin_waiting_broadcast_content", {"target": target})

        elif data.startswith("broadcast_event_"):
            event_id = int(data.split("_")[2])
            await event.edit("لطفا پیام، عکس یا ویدیوی خود را ارسال کنید (در انتهای پیام #پین برای پین کردن اضافه کنید):", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_broadcast_content", {"target": "approved_event", "event_id": event_id})

        elif data == "admin_faq":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, question FROM faqs")
            faqs = c.fetchall()
            conn.close()
            buttons = [[Button.inline("➕ افزودن سوال جدید", b"admin_add_faq")]]
            for faq_id, question in faqs:
                buttons.append([Button.inline(f"✏️ {question[:30]}...", f"admin_edit_faq_{faq_id}")])
                buttons.append([Button.inline(f"🗑️ حذف", f"admin_del_faq_{faq_id}")])
            buttons.append([Button.inline("🏠 بازگشت به منو", b"main_menu")])
            await event.edit("❓ مدیریت سوالات متداول:", buttons=buttons)

        elif data == "admin_add_faq":
            await event.edit("❓ لطفا سوال را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_faq_question")

        elif data.startswith("admin_edit_faq_"):
            faq_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT question, answer FROM faqs WHERE id = ?", (faq_id,))
            row = c.fetchone()
            conn.close()
            if not row:
                await event.answer("❌ سوال یافت نشد!", alert=True)
                return
            question, answer = row
            preview = f"سوال فعلی:\n{question}\n\nپاسخ فعلی:\n{answer}"
            await event.edit(preview + "\n\n✏️ لطفا سوال جدید را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_faq_edit_question", {"faq_id": faq_id})

        elif data.startswith("admin_del_faq_"):
            faq_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("DELETE FROM faqs WHERE id = ?", (faq_id,))
            conn.commit()
            conn.close()
            await event.answer("✅ سوال حذف شد.", alert=True)
            await admin_callback_handler(event)

        elif data == "admin_export_excel":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events WHERE is_active = 1 ORDER BY id DESC")
            events_list = c.fetchall()
            conn.close()
            if not events_list:
                await event.edit("📭 هیچ رویداد فعالی موجود نیست.", buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])
                return
            buttons = []
            for eid, title in events_list:
                buttons.append([Button.inline(title, f"export_excel_event_{eid}")])
            buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit("لطفا رویداد را برای دریافت فایل اکسل انتخاب کنید:", buttons=buttons)

        elif data.startswith("export_excel_event_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT title FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            event_title = row[0] if row else "event"
            df = None
            if pd is not None:
                df = pd.read_sql_query("""
                SELECT 
                    TRIM(u.full_name) AS 'نام و نام خانوادگی',
                    TRIM(u.national_id) AS 'کد ملی',
                    TRIM(u.phone) AS 'شماره تماس'
                FROM registrations r
                JOIN users u ON u.user_id = r.user_id
                WHERE r.status = 'approved' AND r.event_id = ?
                GROUP BY u.full_name, u.national_id, u.phone
                ORDER BY u.full_name COLLATE NOCASE ASC
                """, conn, params=(event_id,))
            else:
                c2 = conn.cursor()
                c2.execute("""
                    SELECT TRIM(u.full_name), TRIM(u.national_id), TRIM(u.phone)
                    FROM registrations r
                    JOIN users u ON u.user_id = r.user_id
                    WHERE r.status = 'approved' AND r.event_id = ?
                    GROUP BY u.full_name, u.national_id, u.phone
                    ORDER BY u.full_name COLLATE NOCASE ASC
                """, (event_id,))
                rows = c2.fetchall()
            conn.close()
            if pd is not None:
                if df.empty:
                    df = pd.DataFrame(columns=['نام و نام خانوادگی', 'کد ملی', 'شماره تماس'])
            output = BytesIO()
            wrote = False
            if pd is not None:
                try:
                    df.to_excel(output, index=False, engine='openpyxl')
                    wrote = True
                except Exception:
                    pass
                if not wrote:
                    try:
                        df.to_excel(output, index=False, engine='xlsxwriter')
                        wrote = True
                    except Exception:
                        pass
            if not wrote:
                # Pure-Python fallback using openpyxl or xlsxwriter
                headers = ['نام و نام خانوادگی', 'کد ملی', 'شماره تماس']
                try:
                    from openpyxl import Workbook
                    wb = Workbook()
                    ws = wb.active
                    ws.title = "approved"
                    ws.append(headers)
                    if pd is not None:
                        for _, row in df.iterrows():
                            ws.append([row.get('نام و نام خانوادگی', ''), row.get('کد ملی', ''), row.get('شماره تماس', '')])
                    else:
                        for r in rows:
                            ws.append([r[0] or '', r[1] or '', r[2] or ''])
                    wb.save(output)
                    wrote = True
                except Exception:
                    pass
            if not wrote:
                try:
                    try:
                        import xlsxwriter as _xlsxwriter  # type: ignore[reportMissingImports]
                    except Exception:
                        _xlsxwriter = None
                    if not _xlsxwriter:
                        raise ImportError('xlsxwriter not available')
                    workbook = _xlsxwriter.Workbook(output)
                    worksheet = workbook.add_worksheet('approved')
                    headers = ['نام و نام خانوادگی', 'کد ملی', 'شماره تماس']
                    for col, header in enumerate(headers):
                        worksheet.write(0, col, header)
                    if pd is not None:
                        for r, (_, row) in enumerate(df.iterrows(), start=1):
                            worksheet.write(r, 0, row.get('نام و نام خانوادگی', ''))
                            worksheet.write(r, 1, row.get('کد ملی', ''))
                            worksheet.write(r, 2, row.get('شماره تماس', ''))
                    else:
                        for r, row in enumerate(rows, start=1):
                            worksheet.write(r, 0, row[0] or '')
                            worksheet.write(r, 1, row[1] or '')
                            worksheet.write(r, 2, row[2] or '')
                    workbook.close()
                    wrote = True
                except Exception as e4:
                    try:
                        await event.answer(f"❌ خطا در تهیه فایل اکسل (xlsx): {str(e4)}", alert=True)
                    except Exception:
                        pass
                    return
                    wrote = True
                except Exception as e4:
                    await event.answer(f"❌ خطا در تهیه فایل اکسل (xlsx): {str(e4)}", alert=True)
                    return
            output.seek(0)
            try:
                from openpyxl import load_workbook
                from openpyxl.utils import get_column_letter
                wb = load_workbook(filename=output)
                for ws in wb.worksheets:
                    try:
                        ws.sheet_view.rightToLeft = True
                    except Exception:
                        pass
                    try:
                        cols = list(ws.columns)
                        for i, col in enumerate(cols, start=1):
                            max_len = 0
                            for cell in col:
                                try:
                                    val = cell.value
                                    if val is None:
                                        continue
                                    s = str(val)
                                except Exception:
                                    s = ''
                                if len(s) > max_len:
                                    max_len = len(s)
                            width = min(50, max(10, int(max_len * 1.2)))
                            col_letter = get_column_letter(i)
                            try:
                                ws.column_dimensions[col_letter].width = width
                            except Exception:
                                pass
                    except Exception:
                        pass
                new_buf = BytesIO()
                wb.save(new_buf)
                new_buf.seek(0)
                output = new_buf
            except Exception:
                # If post-processing fails, continue with original buffer
                output.seek(0)
            try:
                unsafe = ['\\', '/', ':', '*', '?', '"', '<', '>', '|']
                safe_title = event_title
                for ch in unsafe:
                    safe_title = safe_title.replace(ch, ' ')
                safe_title = safe_title.strip()
                file_name = f"approved_{safe_title}.xlsx" if safe_title else "approved_event_users.xlsx"
                try:
                    output.name = file_name
                except Exception:
                    pass
                await client.send_file(event.chat_id, output, caption="📊 لیست تاییدشدگان رویداد", force_document=True)
                await event.answer("✅ فایل اکسل ارسال شد.", alert=True)
            except Exception as e:
                await event.answer(f"❌ خطایی در ارسال فایل اکسل رخ داد: {str(e)}", alert=True)

        elif data == "admin_send_cert":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, title FROM events WHERE is_active = 0")
            events_list = c.fetchall()
            conn.close()
            if not events_list:
                await event.edit("📭 هیچ رویداد آرشیو شده‌ای برای ارسال گواهی وجود ندارد.", buttons=[[Button.inline("🏠 منو", b"main_menu")]])
                return
            buttons = paginate_buttons(events_list, "cert_event", 0, 5)
            await event.edit("📜 لطفا رویداد را انتخاب کنید:", buttons=buttons)

        elif data.startswith("cert_event_"):
            parts = data.split("_")
            if parts[2] == "page":
                page = int(parts[3])
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("SELECT id, title FROM events WHERE is_active = 0")
                events_list = c.fetchall()
                conn.close()
                buttons = paginate_buttons(events_list, "cert_event", page, 5)
                await event.edit("📜 لطفا رویداد را انتخاب کنید:", buttons=buttons)
                return
            event_id = int(parts[2])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("""
                SELECT u.user_id, u.full_name
                FROM registrations r
                JOIN users u ON r.user_id = u.user_id
                WHERE r.event_id = ? AND r.status = 'approved'
            """, (event_id,))
            users = c.fetchall()
            conn.close()
            if not users:
                await event.edit("📭 هیچ کاربر تایید شده‌ای برای این رویداد وجود ندارد.", buttons=[[Button.inline("🔙 بازگشت", b"admin_send_cert")], [Button.inline("🏠 منو", b"main_menu")]])
                return
            buttons = []
            for uid, name in users:
                buttons.append([Button.inline(f"📄 {name}", f"send_cert_to_{uid}_{event_id}")])
            buttons.append([Button.inline("🔙 بازگشت", b"admin_send_cert")])
            buttons.append([Button.inline("🏠 منو", b"main_menu")])
            await event.edit("👤 لطفا کاربر را برای ارسال گواهی انتخاب کنید:", buttons=buttons)

        elif data.startswith("admin_manage_event_regs_"):
            event_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT r.id, r.user_id, COALESCE(u.full_name, '') as full_name, r.status FROM registrations r LEFT JOIN users u ON r.user_id = u.user_id WHERE r.event_id = ? ORDER BY r.id DESC", (event_id,))
            regs = c.fetchall()
            conn.close()
            if not regs:
                await event.edit("📭 هیچ ثبت‌نامی برای این رویداد وجود ندارد.", buttons=[[Button.inline("🔙 بازگشت", f"admin_edit_event_{event_id}")],[Button.inline("🏠 منوی اصلی", b"main_menu")]])
                return
            buttons = []
            for reg_id, uid, full_name, status in regs:
                label = f"#{reg_id} — {full_name or uid} — {status}"
                buttons.append([Button.inline(label, f"admin_view_reg_{reg_id}")])
            buttons.append([Button.inline("🔙 بازگشت", f"admin_edit_event_{event_id}")])
            buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit(f"📋 ثبت‌نام‌های رویداد #{event_id}:", buttons=buttons)

        elif data.startswith("admin_view_reg_"):
            reg_id = int(data.split("_")[3])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT r.user_id, r.event_id, u.full_name, u.national_id, u.phone, u.is_student, u.student_id, r.payment_receipt_file_id, e.title, r.status FROM registrations r JOIN users u ON r.user_id = u.user_id JOIN events e ON r.event_id = e.id WHERE r.id = ?", (reg_id,))
            row = c.fetchone()
            conn.close()
            if not row:
                await event.answer("❌ ثبت‌نام یافت نشد!", alert=True)
                return
            (user_id_t, event_id, name, nid, phone, is_std, std_id, receipt, event_title, status) = row
            status_text = status
            msg = f"📌 ثبت‌نام: #{reg_id}\n👤 نام: {name}\n🆔 کد ملی: {nid}\n📞 تلفن: {phone}\n🎓 {('دانشجو — ' + (std_id or '—')) if is_std else 'غیر دانشجو'}\n📅 رویداد: {event_title}\n📊 وضعیت: {status_text}"
            buttons = [
                [Button.inline("✅ تایید", f"admin_approve_reg_{reg_id}"), Button.inline("❌ رد", f"admin_reject_reg_{reg_id}")],
                [Button.inline("🗑️ حذف ثبت‌نام", f"admin_delete_reg_{reg_id}" )],
                [Button.inline("🔙 بازگشت", f"admin_manage_event_regs_{event_id}"), Button.inline("🏠 منو", b"main_menu")]
            ]
            if receipt and os.path.exists(receipt) and receipt.startswith("uploads" + os.sep):
                try:
                    await client.send_file(event.chat_id, receipt, caption="💳 رسید پرداخت:")
                except Exception:
                    msg += "\n\n⚠️ رسید: [ارسال نشد — خطای فایل]"
            await event.edit(msg, buttons=buttons)

        elif data.startswith("admin_delete_reg_"):
            reg_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, event_id FROM registrations WHERE id = ?", (reg_id,))
            row = c.fetchone()
            if not row:
                conn.close()
                await event.answer("❌ ثبت‌نام یافت نشد!", alert=True)
                return
            user_id_target, event_id = row
            c.execute("DELETE FROM registrations WHERE id = ?", (reg_id,))
            conn.commit()
            conn.close()
            try:
                await client.send_message(user_id_target, f"✅ ثبت‌نام شما برای رویداد #{event_id} حذف شد. اکنون می‌توانید دوباره ثبت‌نام کنید.")
            except Exception:
                pass
            await event.answer("✅ ثبت‌نام حذف شد.", alert=True)
            # attempt to refresh admin view
            try:
                await admin_callback_handler(event)
            except Exception:
                pass

        elif data.startswith("admin_approve_reg_"):
            reg_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, event_id FROM registrations WHERE id = ?", (reg_id,))
            row = c.fetchone()
            if not row:
                conn.close()
                await event.answer("❌ ثبت‌نام یافت نشد!", alert=True)
                return
            user_id_target, event_id = row
            c.execute("UPDATE registrations SET status = 'approved' WHERE id = ?", (reg_id,))
            conn.commit()
            conn.close()
            try:
                await client.send_message(user_id_target, f"✅ ثبت‌نام شما برای رویداد #{event_id} تایید شد. تبریک!")
            except Exception:
                pass
            await event.answer("✅ ثبت‌نام تایید شد.", alert=True)
            try:
                await admin_callback_handler(event)
            except Exception:
                pass

        elif data.startswith("admin_reject_reg_"):
            reg_id = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, event_id FROM registrations WHERE id = ?", (reg_id,))
            row = c.fetchone()
            if not row:
                conn.close()
                await event.answer("❌ ثبت‌نام یافت نشد!", alert=True)
                return
            user_id_target, event_id = row
            c.execute("UPDATE registrations SET status = 'rejected' WHERE id = ?", (reg_id,))
            conn.commit()
            conn.close()
            try:
                await client.send_message(user_id_target, f"❌ ثبت‌نام شما برای رویداد #{event_id} رد شد. در صورت تمایل می‌توانید دوباره ثبت‌نام کنید.")
            except Exception:
                pass
            await event.answer("✅ ثبت‌نام رد شد.", alert=True)
            try:
                await admin_callback_handler(event)
            except Exception:
                pass

        elif data.startswith("send_cert_to_"):
            parts = data.split("_")
            target_user_id = int(parts[3])
            event_id = int(parts[4])
            await event.edit("لطفا فایل PDF گواهی را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_cert_file", {"target_user_id": target_user_id, "event_id": event_id})

        # ---- New admin handlers: ideas / collaborations / donations ----
        elif data == "admin_ideas":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            # show only pending ideas so processed ones disappear from the list
            c.execute("SELECT id, user_id, title, status, created_at FROM ideas WHERE LOWER(TRIM(status)) = 'pending' ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            buttons = []
            if not rows:
                # No ideas yet — show a friendly message but still offer the export button
                buttons.append([Button.inline("📄 دریافت ورد ایده‌ها", b"admin_export_ideas_word")])
                buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
                await event.edit("📭 هیچ ایده‌ای وجود ندارد.", buttons=buttons)
                return
            for iid, uid, title, status, created in rows:
                preview = (title[:40] + '...') if title and len(title) > 40 else (title or 'بدون عنوان')
                buttons.append([Button.inline(f"#{iid} — {preview} — {status}", f"admin_view_idea_{iid}")])
            # show export options
            buttons.append([Button.inline("📄 دریافت ورد ایده‌ها", b"admin_export_ideas_word"), Button.inline("📥 دریافت اکسل ایده‌ها", b"admin_export_ideas_excel")])
            buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit("💡 ایده‌های ارسال شده:", buttons=buttons)

        elif data.startswith("admin_view_idea_"):
            iid = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, title, description, file_path, status, created_at, admin_note, processed_by, processed_at FROM ideas WHERE id = ?", (iid,))
            row = c.fetchone()
            conn.close()
            if not row:
                await event.answer("❌ ایده یافت نشد!", alert=True)
                return
            uid, title, desc, fpath, status, created, admin_note, processed_by, processed_at = row
            msg = f"📌 ایده #{iid}\n👤 کاربر: {uid}\nوضعیت: {status}\nزمان: {created}\n\nعنوان:\n{title or '—'}\n\nتوضیحات:\n{desc or '—'}"
            if admin_note:
                msg += f"\n\n📝 توضیح ادمین:\n{admin_note}"
            if processed_by:
                msg += f"\n\n✅ پردازش شده توسط: {processed_by} در {processed_at}"
            buttons = [
                [Button.inline("✅ تایید", f"admin_approve_idea_{iid}"), Button.inline("❌ رد", f"admin_reject_idea_{iid}")],
                [Button.inline("علامت‌گذاری به عنوان انجام‌شده", f"admin_mark_idea_{iid}")],
                [Button.inline("🔙 بازگشت", b"admin_ideas"), Button.inline("🏠 منو", b"main_menu")]
            ]
            try:
                if fpath and is_safe_upload_path(fpath) and os.path.exists(fpath):
                    await client.send_file(event.chat_id, fpath, caption=msg, buttons=buttons)
                else:
                    await event.edit(msg, buttons=buttons)
            except Exception:
                await event.edit(msg, buttons=buttons)

        elif data.startswith("admin_approve_idea_") or data.startswith("admin_reject_idea_") or data.startswith("admin_mark_idea_"):
            # Prompt admin to enter an explanation that will be sent once to the user
            iid = int(data.split("_")[-1])
            new_status = 'approved' if data.startswith("admin_approve_") else ('rejected' if data.startswith("admin_reject_") else 'handled')
            set_user_state(user_states, user_id, "admin_explain_action", {"target_table": "ideas", "target_id": iid, "new_status": new_status})
            await event.edit("🔎 لطفا توضیحی که می‌خواهید برای کاربر ارسال شود را بنویسید (این پیام تنها یک‌بار ارسال می‌شود):", buttons=CANCEL_BUTTON)

        elif data == "admin_collaborations":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            # show only pending collaboration requests
            c.execute("SELECT id, user_id, full_name, organization, status, created_at FROM collaborations WHERE LOWER(TRIM(status)) = 'pending' ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            buttons = []
            if not rows:
                # No collaborations yet — still allow export (will produce header-only file)
                buttons.append([Button.inline("📄 دریافت ورد همکاری‌ها", b"admin_export_collabs_word")])
                buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
                await event.edit("📭 هیچ درخواستی وجود ندارد.", buttons=buttons)
                return
            for cid, uid, name, org, status, created in rows:
                preview = (name or str(uid))
                buttons.append([Button.inline(f"#{cid} — {preview} — {status}", f"admin_view_collab_{cid}")])
            buttons.append([Button.inline("📄 دریافت ورد همکاری‌ها", b"admin_export_collabs_word"), Button.inline("📥 دریافت اکسل همکاری‌ها", b"admin_export_collabs_excel")])
            buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit("🤝 درخواست‌های همکاری:", buttons=buttons)

        elif data.startswith("admin_view_collab_"):
            cid = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, full_name, organization, proposal, file_path, status, created_at, admin_note, processed_by, processed_at FROM collaborations WHERE id = ?", (cid,))
            row = c.fetchone()
            conn.close()
            if not row:
                await event.answer("❌ درخواست یافت نشد!", alert=True)
                return
            uid, name, org, proposal, fpath, status, created, admin_note, processed_by, processed_at = row
            msg = f"📌 درخواست همکاری #{cid}\n👤 نام: {name}\nسازمان: {org or '—'}\nوضعیت: {status}\nزمان: {created}\n\nپیشنهاد:\n{proposal or '—'}"
            if admin_note:
                msg += f"\n\n📝 توضیح ادمین:\n{admin_note}"
            if processed_by:
                msg += f"\n\n✅ پردازش شده توسط: {processed_by} در {processed_at}"
            buttons = [
                [Button.inline("✅ تایید", f"admin_approve_collab_{cid}"), Button.inline("❌ رد", f"admin_reject_collab_{cid}")],
                [Button.inline("🔙 بازگشت", b"admin_collaborations"), Button.inline("🏠 منو", b"main_menu")]
            ]
            try:
                if fpath and is_safe_upload_path(fpath) and os.path.exists(fpath):
                    await client.send_file(event.chat_id, fpath, caption=msg, buttons=buttons)
                else:
                    await event.edit(msg, buttons=buttons)
            except Exception:
                await event.edit(msg, buttons=buttons)

        elif data.startswith("admin_approve_collab_") or data.startswith("admin_reject_collab_"):
            cid = int(data.split("_")[-1])
            new_status = 'approved' if data.startswith("admin_approve_collab_") else 'rejected'
            set_user_state(user_states, user_id, "admin_explain_action", {"target_table": "collaborations", "target_id": cid, "new_status": new_status})
            await event.edit("🔎 لطفا توضیحی که می‌خواهید برای کاربر ارسال شود را بنویسید (این پیام تنها یک‌بار ارسال می‌شود):", buttons=CANCEL_BUTTON)

        elif data == "admin_donations":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            # show only pending donations for review
            c.execute("SELECT id, user_id, amount, currency, status, created_at FROM donations WHERE LOWER(TRIM(status)) = 'pending' ORDER BY id DESC")
            rows = c.fetchall()
            conn.close()
            buttons = []
            if not rows:
                # No donations yet — keep export available
                buttons.append([Button.inline("📄 دریافت ورد حمایت‌ها", b"admin_export_donations_word")])
                buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
                await event.edit("📭 هیچ حمایتی ثبت نشده است.", buttons=buttons)
                return
            for did, uid, amount, cur, status, created in rows:
                buttons.append([Button.inline(f"#{did} — {amount} {cur} — {status}", f"admin_view_donation_{did}")])
            buttons.append([Button.inline("📄 دریافت ورد حمایت‌ها", b"admin_export_donations_word"), Button.inline("📥 دریافت اکسل حمایت‌ها", b"admin_export_donations_excel")])
            buttons.append([Button.inline("🏠 منوی اصلی", b"main_menu")])
            await event.edit("💰 حمایت‌های ثبت شده:", buttons=buttons)

        elif data.startswith("admin_view_donation_"):
            did = int(data.split("_")[-1])
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id, amount, currency, receipt_file, status, created_at, admin_note, processed_by, processed_at FROM donations WHERE id = ?", (did,))
            row = c.fetchone()
            conn.close()
            if not row:
                await event.answer("❌ مورد یافت نشد!", alert=True)
                return
            uid, amount, cur, receipt, status, created, admin_note, processed_by, processed_at = row
            msg = f"📌 حمایت #{did}\n👤 کاربر: {uid}\nمبلغ: {amount} {cur}\nوضعیت: {status}\نزمان: {created}"
            if admin_note:
                msg += f"\n\n📝 توضیح ادمین:\n{admin_note}"
            if processed_by:
                msg += f"\n\n✅ پردازش شده توسط: {processed_by} در {processed_at}"
            buttons = [
                [Button.inline("✅ تایید و ثبت", f"admin_confirm_donation_{did}"), Button.inline("❌ رد", f"admin_reject_donation_{did}")],
                [Button.inline("🔙 بازگشت", b"admin_donations"), Button.inline("🏠 منو", b"main_menu")]
            ]
            try:
                if receipt and is_safe_upload_path(receipt) and os.path.exists(receipt):
                    await client.send_file(event.chat_id, receipt, caption=msg, buttons=buttons)
                else:
                    await event.edit(msg, buttons=buttons)
            except Exception:
                await event.edit(msg, buttons=buttons)

        elif data.startswith("admin_confirm_donation_") or data.startswith("admin_reject_donation_"):
            did = int(data.split("_")[-1])
            new_status = 'confirmed' if data.startswith("admin_confirm_donation_") else 'rejected'
            set_user_state(user_states, user_id, "admin_explain_action", {"target_table": "donations", "target_id": did, "new_status": new_status})
            await event.edit("🔎 لطفا توضیحی که می‌خواهید برای کاربر ارسال شود را بنویسید (این پیام تنها یک‌بار ارسال می‌شود):", buttons=CANCEL_BUTTON)

        elif data == "admin_tickets":
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT id, user_id, message, created_at FROM tickets WHERE status = 'open'")
            tickets = c.fetchall()
            conn.close()
            buttons = []
            if not tickets:
                # No open tickets — still provide export buttons
                buttons.append([Button.inline("📄 دریافت ورد تیکت‌ها", b"admin_export_tickets_word"), Button.inline("📥 دریافت اکسل تیکت‌ها", b"admin_export_tickets_excel")])
                buttons.append([Button.inline("🏠 منو", b"main_menu")])
                await event.edit("📭 هیچ تیکت بازی وجود ندارد.", buttons=buttons)
                return
            for tid, uid, msg, created in tickets:
                preview = msg[:30] + "..." if len(msg) > 30 else msg
                buttons.append([Button.inline(f"#{tid} — {preview}", f"ticket_reply_{tid}")])
            # show export options
            buttons.append([Button.inline("📄 دریافت ورد تیکت‌ها", b"admin_export_tickets_word"), Button.inline("📥 دریافت اکسل تیکت‌ها", b"admin_export_tickets_excel")])
            buttons.append([Button.inline("🏠 بازگشت به منو", b"main_menu")])
            await event.edit("📬 تیکت‌های باز:", buttons=buttons)

        elif data.startswith("ticket_reply_"):
            ticket_id = int(data.split("_")[2])
            await event.edit("لطفا پاسخ خود را ارسال کنید:", buttons=CANCEL_BUTTON)
            set_user_state(user_states, user_id, "admin_waiting_ticket_reply", {"ticket_id": ticket_id})

    @client.on(events.NewMessage)
    async def admin_registration_flow_handler(event):
        user_id = event.sender_id
        state = get_user_state(user_states, user_id)

        if not utils_is_admin(DB_NAME, user_id):
            return

        if state == "admin_new_event_title":
            title = event.message.text.strip()
            set_user_state(user_states, user_id, "admin_new_event_desc", {"title": title})
            await event.reply("📝 لطفا توضیحات رویداد را ارسال کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_desc":
            desc = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            data["description"] = desc
            set_user_state(user_states, user_id, "admin_new_event_cost_type", data)
            buttons = [
                [Button.inline("رایگان", b"cost_free")],
                [Button.inline("هزینه ثابت", b"cost_fixed")],
                [Button.inline("هزینه متغیر", b"cost_variable")],
                [Button.inline("❌ لغو", b"cancel")]
            ]
            await event.reply("💰 نوع هزینه را انتخاب کنید:", buttons=buttons)

        # Handle admin choosing cost type via callback (cost_free / cost_fixed / cost_variable)

        elif state == "admin_new_event_cost_amount":
            try:
                amount = int(event.message.text.replace(",", "").strip())
                data = get_user_data(user_states, user_id)
                cost_type = data["cost_type"]
                if cost_type == "fixed":
                    data["fixed_cost"] = amount
                    set_user_state(user_states, user_id, "admin_new_event_card", data)
                    await event.reply("💳 لطفا شماره کارت را با فونت مونواسپیس ارسال کنید:", buttons=CANCEL_BUTTON)
                elif cost_type == "variable":
                    data["student_cost"] = amount
                    set_user_state(user_states, user_id, "admin_new_event_non_student_cost", data)
                    await event.reply("🎓 لطفا هزینه برای غیر دانشجو (به تومان) را ارسال کنید:", buttons=CANCEL_BUTTON)
            except ValueError:
                await event.reply("❌ لطفا یک عدد معتبر وارد کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_non_student_cost":
            try:
                amount = int(event.message.text.replace(",", "").strip())
                data = get_user_data(user_states, user_id)
                data["non_student_cost"] = amount
                set_user_state(user_states, user_id, "admin_new_event_card", data)
                await event.reply("💳 لطفا شماره کارت را با فونت مونواسپیس ارسال کنید:", buttons=CANCEL_BUTTON)
            except ValueError:
                await event.reply("❌ لطفا یک عدد معتبر وارد کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_card":
            card = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            data["card_number"] = card
            # DEBUG: log who sets the poster state
            try:
                from log_helper import console_log
                console_log(f"SET_STATE: user={user_id} path=admin_new_event_card -> admin_new_event_poster data_keys={list(data.keys())}")
            except Exception:
                try:
                    print(f"[DEBUG] SET_STATE: user={user_id} path=admin_new_event_card -> admin_new_event_poster data_keys={list(data.keys())}")
                except Exception:
                    pass
            set_user_state(user_states, user_id, "admin_new_event_poster", data)
            await event.reply("🖼️ لطفا پوستر رویداد را ارسال کنید (عکس):", buttons=CANCEL_BUTTON)

        # New states for certificate issuance flow when event is free
        elif state == "admin_new_event_cert_fee":
            # collect certificate issuance fee (integer)
            try:
                amount = int(event.message.text.replace(",", "").strip())
                data = get_user_data(user_states, user_id)
                data["cert_fee"] = amount
                set_user_state(user_states, user_id, "admin_new_event_cert_card", data)
                await event.reply("💳 لطفا شماره کارت برای دریافت هزینه صدور گواهی را ارسال کنید:", buttons=CANCEL_BUTTON)
            except ValueError:
                await event.reply("❌ لطفا یک عدد معتبر وارد کنید (میزان هزینه صدور گواهی به تومان):", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_cert_fee_student":
            try:
                amount = int(event.message.text.replace(",", "").strip())
                data = get_user_data(user_states, user_id)
                data["cert_fee_student"] = amount
                set_user_state(user_states, user_id, "admin_new_event_cert_fee_non_student", data)
                await event.reply("📥 لطفا میزان هزینه صدور گواهی برای غیر دانشجو را به تومان وارد کنید (عدد):", buttons=CANCEL_BUTTON)
            except ValueError:
                await event.reply("❌ لطفا یک عدد معتبر وارد کنید (میزان هزینه صدور گواهی به تومان):", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_cert_fee_non_student":
            try:
                amount = int(event.message.text.replace(",", "").strip())
                data = get_user_data(user_states, user_id)
                data["cert_fee_non_student"] = amount
                set_user_state(user_states, user_id, "admin_new_event_cert_card", data)
                await event.reply("💳 لطفا شماره کارت برای دریافت هزینه صدور گواهی را ارسال کنید:", buttons=CANCEL_BUTTON)
            except ValueError:
                await event.reply("❌ لطفا یک عدد معتبر وارد کنید (میزان هزینه صدور گواهی به تومان):", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_cert_card":
            card = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            data["cert_card_number"] = card
            set_user_state(user_states, user_id, "admin_new_event_cert_card_holder", data)
            await event.reply("🧾 لطفا نام صاحب کارت را وارد کنید (برای صدور گواهی):", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_cert_card_holder":
            holder = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            data["cert_card_holder"] = holder
            # DEBUG: log who sets the poster state
            try:
                from log_helper import console_log
                console_log(f"SET_STATE: user={user_id} path=admin_new_event_cert_card_holder -> admin_new_event_poster data_keys={list(data.keys())}")
            except Exception:
                try:
                    print(f"[DEBUG] SET_STATE: user={user_id} path=admin_new_event_cert_card_holder -> admin_new_event_poster data_keys={list(data.keys())}")
                except Exception:
                    pass
            # proceed to poster upload step same as normal flow
            set_user_state(user_states, user_id, "admin_new_event_poster", data)
            await event.reply("🖼️ لطفا پوستر رویداد را ارسال کنید (عکس):", buttons=CANCEL_BUTTON)

        elif state == "admin_new_event_poster":
            if not event.message.photo:
                await event.reply("❌ لطفا یک عکس ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            f = event.message.photo
            size = getattr(f, 'size', 0) or 0
            if size and size > 5 * 1024 * 1024:
                await event.reply("❌ حجم تصویر پوستر نباید بیش از 5MB باشد.", buttons=CANCEL_BUTTON)
                return
            data = get_user_data(user_states, user_id)
            file_ext = "jpg"
            unique_name = f"poster_{user_id}_{int(time.time())}_{random.randint(1000,9999)}.{file_ext}"
            temp_path = os.path.join("uploads", unique_name)
            await event.message.download_media(file=temp_path)
            data["poster_path"] = temp_path
            title = data["title"]
            desc = data["description"]
            cost_type = data["cost_type"]
            cost_display = "رایگان"
            if cost_type == "fixed":
                cost_display = f"{data['fixed_cost']:,} تومان"
            elif cost_type == "variable":
                cost_display = f"دانشجو: {data['student_cost']:,} — غیر دانشجو: {data['non_student_cost']:,} تومان"
            msg = f"""
📌 *{title}*
{desc or '---'}
            """.strip()
            if cost_type in ["fixed", "variable"]:
                msg += f"\n\n💰 هزینه: {cost_display}"
                msg += f"\n💳 کارت: {data.get('card_number', '---')}"
            # include certificate issuance info for free events if present
            if cost_type == "free":
                # case: single fee
                cert_fee = data.get('cert_fee')
                fee_student = data.get('cert_fee_student')
                fee_non = data.get('cert_fee_non_student')
                if cert_fee is not None:
                    msg += f"\n\n📜 صدور گواهی: بله — هزینه: {cert_fee:,} تومان"
                elif fee_student is not None or fee_non is not None:
                    s_txt = f"دانشجو: {fee_student:,} تومان" if fee_student is not None else "دانشجو: —"
                    ns_txt = f"غیر دانشجو: {fee_non:,} تومان" if fee_non is not None else "غیر دانشجو: —"
                    msg += f"\n\n📜 صدور گواهی: بله — هزینه:\n{s_txt} — {ns_txt}"
                if data.get('cert_card_number') or data.get('cert_card_holder'):
                    msg += f"\n💳 کارت صدور گواهی: {data.get('cert_card_number', '---')}"
                    msg += f"\n👤 صاحب کارت: {data.get('cert_card_holder', '---')}"
            buttons = [
                [Button.inline("✅ تایید و ثبت", b"admin_confirm_event")],
                [Button.inline("✏️ ویرایش", b"admin_edit_event")],
                [Button.inline("❌ لغو", b"cancel")]
            ]
            await client.send_file(
                event.chat_id,
                file=temp_path,
                caption=msg,
                buttons=buttons,
                parse_mode="markdown"
            )
            set_user_state(user_states, user_id, "admin_preview_event", data)
        elif state == "admin_waiting_new_poster":
            if not event.message.photo:
                await event.reply("❌ لطفا یک عکس ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            f = event.message.photo
            size = getattr(f, 'size', 0) or 0
            if size and size > 5 * 1024 * 1024:
                await event.reply("❌ حجم تصویر پوستر نباید بیش از 5MB باشد.", buttons=CANCEL_BUTTON)
                return
            data = get_user_data(user_states, user_id)
            event_id = data["event_id"]
            file_ext = "jpg"
            unique_name = f"poster_{user_id}_{int(time.time())}_{random.randint(1000,9999)}.{file_ext}"
            temp_path = os.path.join("uploads", unique_name)
            await event.message.download_media(file=temp_path)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT poster_file_id FROM events WHERE id = ?", (event_id,))
            row = c.fetchone()
            old_path = row[0] if row else None
            try:
                if old_path and is_safe_upload_path(old_path):
                    os.remove(old_path)
            except Exception:
                pass
            c.execute("UPDATE events SET poster_file_id = ? WHERE id = ?", (temp_path, event_id))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ پوستر رویداد به‌روزرسانی شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_reject_reason":
            reason = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            reg_id = data["reg_id"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE registrations SET status = 'rejected' WHERE id = ?", (reg_id,))
            c.execute("SELECT user_id FROM registrations WHERE id = ?", (reg_id,))
            user_id_target = c.fetchone()[0]
            c.execute("UPDATE users SET reason_if_rejected = ? WHERE user_id = ?", (reason, user_id_target))
            c.execute("UPDATE users SET status = 'rejected' WHERE user_id = ?", (user_id_target,))
            conn.commit()
            conn.close()
            await client.send_message(user_id_target, f"❌ ثبت‌نام شما رد شد.\nدلیل: {reason}")
            clear_user_state(user_states, user_id)
            await event.reply("✅ ثبت‌نام رد شد و دلیل برای کاربر ارسال شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_user_id_to_add":
            try:
                target_id = int(event.message.text.strip())
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                c.execute("INSERT OR IGNORE INTO admins (user_id, added_by) VALUES (?, ?)", (target_id, user_id))
                conn.commit()
                conn.close()
                clear_user_state(user_states, user_id)
                await event.reply("✅ کاربر به عنوان ادمین اضافه شد.", buttons=get_admin_main_menu())
            except ValueError:
                await event.reply("❌ لطفا یک آیدی عددی معتبر وارد کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_waiting_broadcast_content":
            data = get_user_data(user_states, user_id)
            target = data["target"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            if target == "all":
                c.execute("SELECT user_id FROM users")
            elif target == "approved":
                c.execute("""
                    SELECT DISTINCT r.user_id
                    FROM registrations r
                    WHERE r.status = 'approved'
                """)
            elif target == "rejected":
                c.execute("""
                    SELECT DISTINCT u.user_id
                    FROM users u
                    JOIN registrations r ON r.user_id = u.user_id
                    WHERE r.status = 'rejected'
                """)
            elif target == "approved_event":
                event_id = int(data.get("event_id"))
                c.execute("""
                    SELECT DISTINCT r.user_id
                    FROM registrations r
                    WHERE LOWER(TRIM(r.status)) = 'approved' AND r.event_id = ?
                """, (event_id,))
            elif target == "rejected_event":
                event_id = int(data.get("event_id"))
                c.execute("""
                    SELECT DISTINCT r.user_id
                    FROM registrations r
                    WHERE LOWER(TRIM(r.status)) = 'rejected' AND r.event_id = ?
                """, (event_id,))
            recipients = [row[0] for row in c.fetchall()]
            conn.close()
            if not recipients:
                await event.reply("📭 هیچ کاربری در این گروه وجود ندارد. اگر هدف 'تاییدشدگان یک رویداد' است، ابتدا از منوی 'ارسال همگانی' رویداد را انتخاب کنید یا وضعیت ثبت‌نام‌ها را بررسی کنید.", buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])
                clear_user_state(user_states, user_id)
                return
            should_pin = False
            text = event.message.text or ""
            if text.endswith("#پین"):
                should_pin = True
                text = text[:-4].strip()
            total = len(recipients)
            sent_count = 0
            failed_count = 0
            status_msg = await event.reply(f"📤 در حال ارسال به {total} کاربر...")
            from utils import send_with_rate_limit
            for i, uid in enumerate(recipients, 1):
                try:
                    if event.message.file:
                        res = await send_with_rate_limit(client, uid, text=text or None, file=event.message.file)
                    else:
                        res = await send_with_rate_limit(client, uid, text=text)
                    if should_pin and res:
                        try:
                            await client.pin_message(uid, res, notify=False)
                        except Exception:
                            pass
                    if res:
                        sent_count += 1
                    else:
                        failed_count += 1
                except Exception as e:
                    failed_count += 1
                if i % 10 == 0 or i == total:
                    try:
                        await status_msg.edit(f"📤 ارسال به {total} کاربر...\n✅ ارسال شده: {sent_count}\n❌ خطا: {failed_count}")
                    except:
                        pass
                await asyncio.sleep(0.12)
            clear_user_state(user_states, user_id)
            summary = f"✅ ارسال همگانی به پایان رسید.\n\n📬 کل کاربران: {total}\n✅ دریافت کردند: {sent_count}\n❌ ارسال نشد: {failed_count}"
            try:
                await status_msg.edit(summary, buttons=[[Button.inline("🏠 منوی اصلی", b"main_menu")]])
            except:
                pass

        elif state == "admin_waiting_faq_question":
            question = event.message.text.strip()
            set_user_state(user_states, user_id, "admin_waiting_faq_answer", {"question": question})
            await event.reply("📝 لطفا پاسخ سوال را ارسال کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_waiting_faq_answer":
            answer = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            question = data["question"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("INSERT INTO faqs (question, answer) VALUES (?, ?)", (question, answer))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ سوال متداول اضافه شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_faq_edit_question":
            new_question = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            faq_id = data["faq_id"]
            set_user_state(user_states, user_id, "admin_waiting_faq_edit_answer", {"faq_id": faq_id, "question": new_question})
            await event.reply("📝 لطفا پاسخ جدید را ارسال کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_waiting_faq_edit_answer":
            new_answer = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            faq_id = data["faq_id"]
            new_question = data["question"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE faqs SET question = ?, answer = ? WHERE id = ?", (new_question, new_answer, faq_id))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ سوال ویرایش شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_deadline":
            text = (event.message.text or "").strip()
            data = get_user_data(user_states, user_id)
            event_id = data["event_id"]
            from utils import parse_jalali_to_epoch
            ts = parse_jalali_to_epoch(text)
            if not ts:
                await event.reply("❌ فرمت نادرست است. نمونه معتبر: 1403/07/01 18:30", buttons=CANCEL_BUTTON)
                return
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE events SET end_at_ts = ?, end_set_by = ? WHERE id = ?", (int(ts), user_id, event_id))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ مهلت ثبت‌نام تنظیم شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_cert_file":
            if not event.message.file:
                await event.reply("❌ لطفا یک فایل PDF ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            f = event.message.file
            ext = (f.ext or "").lower().lstrip('.')
            size = getattr(f, 'size', 0) or 0
            if ext != 'pdf' or size > 30 * 1024 * 1024:
                await event.reply("❌ فقط PDF تا حجم15MB مجاز است.", buttons=CANCEL_BUTTON)
                return
            data = get_user_data(user_states, user_id)
            target_user_id = data["target_user_id"]
            event_id = data["event_id"]
            file_ext = event.message.file.ext or "pdf"
            unique_name = f"cert_{target_user_id}_{event_id}_{int(time.time())}_{random.randint(1000,9999)}.{file_ext}"
            file_path = os.path.join("uploads", unique_name)
            await event.message.download_media(file=file_path)
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("INSERT INTO certificates (user_id, event_id, file_id, sent_by_admin) VALUES (?, ?, ?, ?)",
                      (target_user_id, event_id, file_path, user_id))
            conn.commit()
            conn.close()
            try:
                await client.send_file(target_user_id, file_path, caption="📜 گواهی شما آماده است!")
                await event.reply("✅ گواهی برای کاربر ارسال شد و در سیستم ثبت گردید.", buttons=get_admin_main_menu())
            except:
                await event.reply("⚠️ گواهی در سیستم ثبت شد اما ارسال به کاربر با خطا مواجه شد.", buttons=get_admin_main_menu())
            clear_user_state(user_states, user_id)

        elif state == "admin_waiting_donation_card":
            card_text = (event.message.text or "").strip()
            # simple sanitize: keep digits and spaces
            cleaned = ''.join([c for c in card_text if c.isdigit() or c.isspace()])
            if not cleaned:
                await event.reply("❌ شماره کارت معتبر نیست. لطفا فقط اعداد را ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            try:
                # save card number temporarily and ask for holder name
                set_setting('donation_card_number', cleaned)
            except Exception:
                pass
            # ask for card holder name
            set_user_state(user_states, user_id, "admin_waiting_donation_holder", {"card_number": cleaned})
            await event.reply("🧾 لطفا نام صاحب کارت (مثلا: نام و نام خانوادگی) را ارسال کنید:", buttons=CANCEL_BUTTON)

        elif state == "admin_waiting_donation_desc":
            desc = event.message.text or ''
            try:
                set_setting('donation_description', desc)
            except Exception:
                pass
            clear_user_state(user_states, user_id)
            await event.reply("✅ توضیحات حمایت ذخیره شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_donation_holder":
            holder = (event.message.text or '').strip()
            if not holder:
                await event.reply("❌ نام صاحب کارت معتبر نیست. لطفا دوباره ارسال کنید:", buttons=CANCEL_BUTTON)
                return
            try:
                set_setting('donation_card_holder', holder)
            except Exception:
                pass
            clear_user_state(user_states, user_id)
            await event.reply(f"✅ اطلاعات کارت حمایت ذخیره شد:\nشماره: {get_setting('donation_card_number','—')}\nصاحب کارت: {holder}", buttons=get_admin_main_menu())

        elif state == "admin_waiting_ticket_reply":
            reply_text = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            ticket_id = data["ticket_id"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("SELECT user_id FROM tickets WHERE id = ?", (ticket_id,))
            user_id_target = c.fetchone()[0]
            c.execute("UPDATE tickets SET status = 'closed', admin_reply = ?, replied_at = datetime('now', 'localtime') WHERE id = ?", (reply_text, ticket_id))
            conn.commit()
            conn.close()
            await client.send_message(user_id_target, f"📬 پاسخ ادمین به تیکت شما:\n\n{reply_text}")
            clear_user_state(user_states, user_id)
            await event.reply("✅ پاسخ ارسال شد و تیکت بسته شد.", buttons=get_admin_main_menu())

        elif state == "admin_explain_action":
            # Admin entered an explanation for an action on idea/collab/donation
            data = get_user_data(user_states, user_id) or {}
            target_table = data.get("target_table")
            target_id = data.get("target_id")
            new_status = data.get("new_status")
            msg_text = (event.message.text or "").strip()
            file_obj = getattr(event.message, 'file', None)

            if not target_table or not target_id or not new_status:
                clear_user_state(user_states, user_id)
                await event.reply("❌ خطا: اطلاعات عملیات پیدا نشد. دوباره تلاش کنید.", buttons=get_admin_main_menu())
                return

            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            try:
                # fetch target user and current status
                if target_table == 'ideas':
                    c.execute("SELECT user_id, LOWER(TRIM(status)) FROM ideas WHERE id = ?", (target_id,))
                elif target_table == 'collaborations':
                    c.execute("SELECT user_id, LOWER(TRIM(status)) FROM collaborations WHERE id = ?", (target_id,))
                elif target_table == 'donations':
                    c.execute("SELECT user_id, LOWER(TRIM(status)) FROM donations WHERE id = ?", (target_id,))
                else:
                    c.execute("SELECT NULL, NULL")
                row = c.fetchone()
                target_uid = row[0] if row else None
                current_status = row[1] if row and len(row) > 1 else None

                # only update if still pending to avoid double-processing
                if current_status and current_status != 'pending':
                    conn.close()
                    clear_user_state(user_states, user_id)
                    await event.reply(f"⚠️ این مورد قبلاً پردازش شده (وضعیت فعلی: {current_status}). عملیات لغو شد.", buttons=get_admin_main_menu())
                    return

                # update status and admin metadata in DB
                if target_table in ('ideas', 'collaborations', 'donations'):
                    q = f"UPDATE {target_table} SET status = ?, admin_note = ?, processed_by = ?, processed_at = datetime('now','localtime') WHERE id = ?"
                    c.execute(q, (new_status, msg_text or None, user_id, target_id))
                # record admin action (best-effort)
                try:
                    c.execute('''CREATE TABLE IF NOT EXISTS admin_actions (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        admin_id INTEGER NOT NULL,
                        action TEXT NOT NULL,
                        target_table TEXT,
                        target_id INTEGER,
                        note TEXT,
                        created_at TEXT DEFAULT (datetime('now','localtime'))
                    )''')
                    c.execute('INSERT INTO admin_actions (admin_id, action, target_table, target_id, note) VALUES (?, ?, ?, ?, ?)',
                              (user_id, f'set_status_{new_status}', target_table, target_id, msg_text[:2000] if msg_text else None))
                except Exception:
                    pass

                conn.commit()
            except Exception:
                conn.rollback()
                conn.close()
                clear_user_state(user_states, user_id)
                await event.reply("❌ خطا در به‌روزرسانی دیتابیس.", buttons=get_admin_main_menu())
                return
            conn.close()

            sent_ok = False
            if target_uid:
                try:
                    from utils import send_with_rate_limit
                    # prepare a templated message similar to registration notifications
                    status_label = ''
                    if new_status == 'approved' or new_status == 'confirmed':
                        status_label = '✅ تایید شد'
                    elif new_status == 'rejected':
                        status_label = '❌ رد شد'
                    elif new_status == 'handled':
                        status_label = '🔔 علامت‌گذاری شد (انجام‌شده)'
                    else:
                        status_label = new_status

                    header = ''
                    if target_table == 'ideas':
                        header = f"💡 وضعیت ایده شما (#{target_id}): {status_label}\n\n"
                    elif target_table == 'collaborations':
                        header = f"🤝 وضعیت درخواست همکاری شما (#{target_id}): {status_label}\n\n"
                    elif target_table == 'donations':
                        header = f"💰 وضعیت حمایت مالی شما (#{target_id}): {status_label}\n\n"

                    body = msg_text or ''
                    # fallback messages if admin didn't provide text
                    if not body:
                        if new_status == 'approved' or new_status == 'confirmed':
                            body = 'درخواست شما بررسی و تایید شد. از همراهی شما سپاسگزاریم.'
                        elif new_status == 'rejected':
                            body = 'متاسفانه درخواست شما مورد تایید قرار نگرفت.'
                        else:
                            body = 'وضعیت درخواست شما به‌روز شد.'

                    final_text = header + body

                    # send with optional file attachment
                    if file_obj:
                        res = await send_with_rate_limit(client, target_uid, text=final_text, file=file_obj)
                    else:
                        res = await send_with_rate_limit(client, target_uid, text=final_text)
                    sent_ok = bool(res)
                except Exception:
                    sent_ok = False

            clear_user_state(user_states, user_id)
            # confirm to admin and return to admin menu
            if target_uid:
                if sent_ok:
                    await event.reply("✅ وضعیت مورد با موفقیت به‌روزرسانی شد و پیام توضیح برای کاربر ارسال شد.", buttons=get_admin_main_menu())
                else:
                    await event.reply("⚠️ وضعیت مورد به‌روزرسانی شد اما ارسال پیام به کاربر با خطا مواجه شد. لطفا بررسی کنید.", buttons=get_admin_main_menu())
            else:
                await event.reply("⚠️ وضعیت مورد به‌روزرسانی شد اما کاربر مقصد پیدا نشد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_resource":
            clear_user_state(user_states, user_id)
            await event.reply("ℹ️ بخش منابع حذف شده است.", buttons=get_admin_main_menu())

        elif state == "edit_event_title":
            new_title = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            event_id = data["event_id"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE events SET title = ? WHERE id = ?", (new_title, event_id))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ عنوان با موفقیت ویرایش شد.", buttons=get_admin_main_menu())

        elif state == "edit_event_desc":
            new_desc = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            event_id = data["event_id"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE events SET description = ? WHERE id = ?", (new_desc, event_id))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ توضیحات با موفقیت ویرایش شد.", buttons=get_admin_main_menu())

        elif state == "edit_event_card":
            new_card = event.message.text.strip()
            data = get_user_data(user_states, user_id)
            event_id = data["event_id"]
            conn = sqlite3.connect(DB_NAME)
            c = conn.cursor()
            c.execute("UPDATE events SET card_number = ? WHERE id = ?", (new_card, event_id))
            conn.commit()
            conn.close()
            clear_user_state(user_states, user_id)
            await event.reply("✅ شماره کارت با موفقیت ویرایش شد.", buttons=get_admin_main_menu())

        elif state == "admin_waiting_reports":
            msg_text = event.message.text or ""
            if msg_text.startswith("/done"):
                data = get_user_data(user_states, user_id)
                event_id = data["event_id"]
                msg_ids = ",".join(map(str, data.get("message_ids", [])))
                payloads = data.get("payloads", [])
                payloads_json = json.dumps(payloads, ensure_ascii=False)
                conn = sqlite3.connect(DB_NAME)
                c = conn.cursor()
                try:
                    c.execute("ALTER TABLE events ADD COLUMN report_payloads TEXT")
                except Exception:
                    pass
                c.execute("UPDATE events SET report_message_ids = ?, report_payloads = ? WHERE id = ?", (msg_ids, payloads_json, event_id))
                conn.commit()
                conn.close()
                clear_user_state(user_states, user_id)
                await event.reply("✅ گزارش کار تنظیم شد.", buttons=get_admin_main_menu())
            else:
                data = get_user_data(user_states, user_id)
                fwd = getattr(event.message, "fwd_from", None)
                forwarded_from_channel = False
                if fwd and getattr(fwd, "channel_post", None):
                    forwarded_from_channel = True
                    data.setdefault("message_ids", []).append(fwd.channel_post)
                    try:
                        channel_id = None
                        if hasattr(fwd, "from_id") and fwd.from_id:
                            channel_id = getattr(fwd.from_id, "channel_id", None) or getattr(fwd.from_id, "chat_id", None)
                        if channel_id:
                            fwd_payload = {
                                "type": "forward",
                                "channel_id": int(channel_id),
                                "message_id": int(fwd.channel_post)
                            }
                            if msg_text:
                                try:
                                    fwd_payload["caption"] = msg_text
                                except Exception:
                                    pass
                            data.setdefault("payloads", []).append(fwd_payload)
                    except Exception:
                        pass

                if event.message.file and not forwarded_from_channel:
                    file_ext = event.message.file.ext or "dat"
                    unique_name = f"report_{user_id}_{int(time.time())}_{random.randint(1000,9999)}.{file_ext}"
                    file_path = os.path.join("uploads", unique_name)
                    await event.message.download_media(file=file_path)
                    data.setdefault("payloads", []).append({
                        "type": "file",
                        "path": file_path,
                        "caption": msg_text
                    })
                elif msg_text and not forwarded_from_channel:
                    data.setdefault("payloads", []).append({
                        "type": "text",
                        "text": msg_text
                    })
                set_user_state(user_states, user_id, "admin_waiting_reports", data)
                await event.reply("✅ پیام گزارش افزوده شد. اگر تمام شد، /done را بفرستید.")


    @client.on(events.CallbackQuery(pattern=b"cost_(free|fixed|variable)"))
    async def cost_type_handler(event):
        user_id = event.sender_id
        if not utils_is_admin(DB_NAME, user_id):
            return
        cost_type = event.data.decode('utf-8').split("_")[1]
        # Only handle fixed/variable here. "free" is handled in the main callback handler
        if cost_type == "free":
            # Do nothing here; the primary handler will present certificate options
            return
        data = get_user_data(user_states, user_id) or {}
        data["cost_type"] = cost_type
        set_user_state(user_states, user_id, "admin_new_event_cost_amount", data)
        if cost_type == "fixed":
            await event.edit("💰 لطفا مبلغ ثابت (به تومان) را ارسال کنید:", buttons=CANCEL_BUTTON)
        elif cost_type == "variable":
            await event.edit("🎓 لطفا هزینه برای دانشجو (به تومان) را ارسال کنید:", buttons=CANCEL_BUTTON)

    @client.on(events.CallbackQuery(data=b"admin_confirm_event"))
    async def confirm_event_handler(event):
        user_id = event.sender_id
        if not utils_is_admin(DB_NAME, user_id):
            return
        data = get_user_data(user_states, user_id)
        if "poster_path" not in data:
            await event.answer("❌ پوستر ذخیره نشده است!", alert=True)
            return
        poster_path = data["poster_path"]
        poster_file_id = poster_path
        conn = sqlite3.connect(DB_NAME)
        c = conn.cursor()
        # store certificate-related fields if provided
        c.execute("""
            INSERT INTO events (
                title, description, cost_type, fixed_cost, student_cost, non_student_cost,
                card_number, poster_file_id, created_by, cert_fee, cert_fee_student, cert_fee_non_student, cert_card_number, cert_card_holder
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            data["title"], data["description"], data["cost_type"],
            data.get("fixed_cost", 0), data.get("student_cost", 0), data.get("non_student_cost", 0),
            data.get("card_number"), poster_file_id, user_id,
            data.get("cert_fee"), data.get("cert_fee_student"), data.get("cert_fee_non_student"), data.get("cert_card_number"), data.get("cert_card_holder")
        ))
        conn.commit()
        conn.close()
        clear_user_state(user_states, user_id)
        await event.edit("✅ رویداد با موفقیت ثبت شد!", buttons=get_admin_main_menu())

    def get_admin_main_menu():
        from utils import get_main_menu_buttons

        return get_main_menu_buttons(True)